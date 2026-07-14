"""
Εξαγωγή βιβλίου εκρηκτικών.
2 σελίδες landscape A4:
  Σελίδα 1: Αγορές / Επιστροφές
  Σελίδα 2: Καταναλώσεις
"""
import io, os, sys, glob, re
from datetime import datetime
from collections import OrderedDict

def find_font(names):
    search_dirs = []
    # Bundled fonts πρώτα (EXPVAULT_FONTS_DIR από main.js)
    bundled = os.environ.get('EXPVAULT_FONTS_DIR')
    if bundled and os.path.isdir(bundled):
        search_dirs.append(bundled)
    search_dirs += [
        '/usr/share/fonts', '/usr/share/fonts/TTF',
        '/usr/local/share/fonts',
        os.path.expanduser('~/.fonts'),
        os.path.expanduser('~/.local/share/fonts'),
        os.path.expanduser('~/.cargo'),
    ]
    if sys.platform == 'win32':
        search_dirs.extend([
            r'C:\Windows\Fonts',
            os.path.expanduser(r'~\AppData\Local\Microsoft\Windows\Fonts'),
        ])
    for name in names:
        for d in search_dirs:
            matches = glob.glob(f'{d}/**/{name}', recursive=True)
            if matches: return matches[0]
            direct = os.path.join(d, name)
            if os.path.exists(direct): return direct
    return None

FONT_SETS = {
    'iosevka':    ['Iosevka-Regular.ttf', 'JetBrainsMono-Regular.ttf', 'LiberationMono-Regular.ttf', 'cour.ttf'],
    'jetbrains':  ['JetBrainsMono-Regular.ttf', 'Iosevka-Regular.ttf', 'LiberationMono-Regular.ttf', 'cour.ttf'],
    'liberation': ['LiberationMono-Regular.ttf', 'FreeMono.ttf', 'cour.ttf'],
}
FONT_SETS_BOLD = {
    'iosevka':    ['Iosevka-Bold.ttf', 'JetBrainsMono-Bold.ttf', 'LiberationMono-Bold.ttf', 'courbd.ttf'],
    'jetbrains':  ['JetBrainsMono-Bold.ttf', 'Iosevka-Bold.ttf', 'LiberationMono-Bold.ttf', 'courbd.ttf'],
    'liberation': ['LiberationMono-Bold.ttf', 'FreeMonoBold.ttf', 'courbd.ttf'],
}

FONT_REGULAR = find_font(FONT_SETS['iosevka'])
FONT_BOLD    = find_font(FONT_SETS_BOLD['iosevka'])

def register_fonts(font_name='iosevka'):
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    reg  = find_font(FONT_SETS.get(font_name, FONT_SETS['iosevka']))
    bold = find_font(FONT_SETS_BOLD.get(font_name, FONT_SETS_BOLD['iosevka']))
    fname = f'Font_{font_name}'
    fbold = f'Font_{font_name}_Bold'
    try:
        pdfmetrics.getFont(fname)
    except:
        if reg:
            pdfmetrics.registerFont(TTFont(fname, reg))
            pdfmetrics.registerFont(TTFont(fbold, bold or reg))
    return fname, fbold

def fmt_date(s):
    if not s: return ''
    try:
        p = s.split('-')
        if len(p)==3 and len(p[0])==4: return f"{p[2]}/{p[1]}/{p[0]}"
    except: pass
    return s

def fmt_num(val):
    if val is None or val == 0: return ''
    return f"{val:,.2f}".replace(',','X').replace('.',',').replace('X','.')

# Σειρά εμφάνισης υλικών στο export
YLIKA_ORDER_IDS = [1, 4, 3, 2, 5, 10, 9, 33]

def _is_nonel(ydata):
    grp = ydata.get('export_group') or ''
    return grp.startswith('NONEL ')

def _nonel_sub(ydata):
    """Επιστρέφει 'SNAPLINE', 'LP' ή 'UNIDET' από το export_group."""
    grp = ydata.get('export_group') or ''
    parts = grp.split()
    return parts[1] if len(parts) > 1 else ''

def _sort_key(yid, ydata):
    if _is_nonel(ydata):
        return 999
    try:
        return YLIKA_ORDER_IDS.index(yid)
    except ValueError:
        return 500

def get_ylika_order(kiniseis, nonel_mode='detail'):
    from database import get_all_ylika
    all_ylika = {y['id']: y for y in get_all_ylika()}

    # Fixed σειρά: ANFO, POLADYN 65, POLADYN 38, EM-EX, ΒΡΑΔΥΚΑΥΣΤΗ, ΑΚΑΡΙΑΙΑ, ΚΟΙΝΟΙ, ΗΛΕΚΤΡΙΚΟΙ, NONEL
    FIXED = [1, 4, 3, 2, 5, 10, 9, 33]
    seen_ids = set(k['yliko_id'] for k in kiniseis)

    if nonel_mode == 'detail':
        result = []
        added = set()
        for yid in FIXED:
            if yid in all_ylika:
                y = all_ylika[yid]
                result.append((yid, (y['onoma'], y['monada_metrisis'])))
                added.add(yid)
        for k in kiniseis:
            yid = k['yliko_id']
            if yid not in added and not _is_nonel(all_ylika.get(yid, {})):
                result.append((yid, (k['yliko_onoma'], k['monada_metrisis'])))
                added.add(yid)
        nonel_added = set()
        for k in kiniseis:
            yid = k['yliko_id']
            if _is_nonel(all_ylika.get(yid, {})) and yid not in nonel_added:
                result.append((yid, (k['yliko_onoma'], k['monada_metrisis'])))
                nonel_added.add(yid)
        return result

    seen_yids = OrderedDict()
    for k in kiniseis:
        yid = k['yliko_id']
        if yid not in seen_yids:
            seen_yids[yid] = all_ylika.get(yid, {})

    result = []
    added = set()
    nonel_added = set()

    for yid in FIXED:
        if yid not in seen_yids or yid in added:
            continue
        ydata = seen_yids[yid]
        result.append((yid, (ydata.get('onoma', ''), ydata.get('monada_metrisis', ''))))
        added.add(yid)

    for yid, ydata in seen_yids.items():
        if yid in added or _is_nonel(ydata):
            continue
        result.append((yid, (ydata.get('onoma', ''), ydata.get('monada_metrisis', ''))))
        added.add(yid)

    for yid, ydata in seen_yids.items():
        if not _is_nonel(ydata):
            continue
        if nonel_mode == 'total':
            if 'NONEL_TOTAL' not in nonel_added:
                result.append(('NONEL_TOTAL', ('NONEL ΣΥΝΟΛΟ', 'Τεμ')))
                nonel_added.add('NONEL_TOTAL')
        elif nonel_mode == 'grouped':
            for s in ['SNAPLINE', 'UNIDET', 'LP']:
                key = 'NONEL_' + s
                if key not in nonel_added:
                    result.append((key, ('NONEL ' + s, 'Τεμ')))
                    nonel_added.add(key)
        elif nonel_mode == 'subgroup':
            gkey = nonel_group_key(ydata)
            vkey = 'NONEL_DLY_{}_{}'.format(*gkey)
            if vkey not in nonel_added:
                onoma = ydata.get('onoma', '')
                display = re.sub(r'\s*\([^)]*\)', '', onoma).strip().replace(',', '').replace('  ', ' ')
                result.append((vkey, (display, 'Τεμ')))
                nonel_added.add(vkey)

    return result

def get_nonel_sums(row_ylika, all_ylika, nonel_mode):
    """Υπολογίζει αθροίσματα NONEL για grouped/total mode."""
    sums = {'SNAPLINE': 0, 'UNIDET': 0, 'LP': 0}
    for yid, pos in row_ylika.items():
        y = all_ylika.get(yid, {})
        if _is_nonel(y):
            sub = _nonel_sub(y)
            if sub in sums:
                sums[sub] += pos or 0
    sums['TOTAL'] = sum(sums.values())
    return sums

def nonel_group_key(ydata):
    """Εξάγει (subgroup, type_key) για ομαδοποίηση NONEL ανά τύπο/καθυστέρηση.
    π.χ. export_group='NONEL SNAPLINE', onoma='NONEL SNAPLINE, SL17 (4.8M)' → ('SNAPLINE', 'SL17')
    """
    sub   = _nonel_sub(ydata)
    onoma = ydata.get('onoma', '')
    m = re.search(r'([A-Z][A-Z0-9.]*)\s*\(', onoma)
    return (sub, m.group(1) if m else onoma)

# ── Ομαδοποιημένο header (PDF + Excel, nonel_mode='detail') ──────────────────
# Κοινή λογική: αντί για ένα flat "NONEL SNAPLINE, SL17 (6.0M) (Τεμ)" ανά στήλη,
# 3 επίπεδα: NONEL -> SNAPLINE/UNIDET U500/LP -> SL17/SL42/.../μήκος. Ίδιο και
# για POLADYN/ΘΡΥΑΛ./ΠΥΡΟΚΡ. (γονέας + leaves).
GROUPED_PARENT_LEAVES = {
    'POLADYN':     ('POLADYN', ('65X500', '38X380')),
    'THRYALLIDES': ('ΘΡΥΑΛ.',  ('ΒΡΑΔΥ', 'ΑΚΑΡ')),
    'KAPSYLIA':    ('ΠΥΡΟΚΡ.', ('ΚΟΙΝΟΙ', 'ΗΛΕΚΤΡ')),
}
GROUPED_SHORT_LABEL = {1: ('AN-FO',), 2: ('EM-EX', 'LC-30')}
GROUPED_NONEL_SUB_ORDER = ['SNAPLINE', 'UNIDET', 'LP']
GROUPED_SUB_LABEL = {'SNAPLINE': 'SNAPLINE', 'UNIDET': 'UNIDET U500', 'LP': 'LP'}

def _nonel_code_length(onoma):
    m = re.search(r'([A-Z][A-Z0-9.]*)\s*\(([^)]*)\)', onoma)
    return (m.group(1) if m else onoma, m.group(2).strip() if m else '')

def grouped_header_plan(ylika_order, all_y):
    """Ξαναταξινομεί το ylika_order ώστε τα NONEL να μαζεύονται ανά (subgroup,
    code, length) και επιστρέφει (ylika_order, n_fixed, n_nonel, nonel_meta)
    όπου nonel_meta: yid -> (sub, code, length). Χρησιμοποιείται μόνο σε
    nonel_mode='detail' — τα άλλα modes έχουν ήδη λίγες, σύντομες στήλες.
    """
    nonel_items, other_items = [], []
    for entry in ylika_order:
        yid, (on, mo) = entry
        ydata = all_y.get(yid, {}) if isinstance(yid, int) else {}
        if isinstance(yid, int) and _is_nonel(ydata):
            sub = _nonel_sub(ydata)
            code, length = _nonel_code_length(on)
            nonel_items.append((entry, sub, code, length))
        else:
            other_items.append(entry)
    nonel_items.sort(key=lambda it: (
        GROUPED_NONEL_SUB_ORDER.index(it[1]) if it[1] in GROUPED_NONEL_SUB_ORDER else 99, it[2], it[3]))
    nonel_meta = {entry[0]: (sub, code, length) for entry, sub, code, length in nonel_items}
    new_order = other_items + [it[0] for it in nonel_items]
    return new_order, len(other_items), len(nonel_items), nonel_meta

def build_book_rows(kiniseis):
    agores     = OrderedDict()
    epistrofes = []
    katanaliseis = []

    for k in kiniseis:
        yid   = k['yliko_id']
        tipos = k['tipos']
        parst = k.get('arithmos_parstatikos') or ''
        imer  = k['imerominia']
        adeia = k.get('arithmos_adeias') or ''
        ekd   = k.get('syntomografia_ekdousas') or k.get('ekdousa_archi') or ''
        prom  = k.get('promitheftis_syntomografia') or k.get('promitheftis_onoma') or ''
        par   = k.get('paratirishis') or ''

        if tipos == 'ΕΙΣΑΓΩΓΗ':
            key = (imer, parst, adeia, prom)
            if key not in agores:
                agores[key] = {
                    'type':'agora', 'aa':0,
                    'imerominia':imer, 'parstatiko':parst,
                    'adeia':adeia, 'ekdousa':ekd, 'promitheftis':prom,
                    'ylika':{}, 'paratirishis':par,
                    'auxon': k.get('auxon_arithmos', 0)
                }
            agores[key]['ylika'][yid] = agores[key]['ylika'].get(yid,0) + k['posotita']

        elif tipos in ('ΕΠΙΣΤΡΟΦΗ', 'ΕΞΑΓΩΓΗ') and parst:
            found = next((e for e in epistrofes
                         if e['parstatiko']==parst and e['imerominia']==imer), None)
            if not found:
                found = {
                    'type':'epistrofi', 'aa':None,
                    'imerominia':imer, 'parstatiko':parst,
                    'adeia':adeia, 'ekdousa':ekd, 'promitheftis':prom,
                    'ylika':{}, 'paratirishis':'ΕΠΙΣΤΡΟΦΗ',
                    'auxon': k.get('auxon_arithmos', 0),
                    'agora_ref': k.get('agora_ref') or ''
                }
                epistrofes.append(found)
            found['ylika'][yid] = found['ylika'].get(yid,0) + k['posotita']

        else:
            # ΚΑΤΑΝΑΛΩΣΗ — ομαδοποίηση ανά parstatiko (αν υπάρχει) ή imerominia
            kat_key = parst if parst else imer
            found = next((c for c in katanaliseis if c['kat_key']==kat_key), None)
            if not found:
                found = {
                    'type':'katanalosi', 'aa':None,
                    'imerominia':imer, 'parstatiko':parst or '',
                    'kat_key': kat_key,
                    'adeia':'', 'ekdousa':'', 'promitheftis':'',
                    'ylika':{}, 'paratirishis':par
                }
                katanaliseis.append(found)
            found['ylika'][yid] = found['ylika'].get(yid,0) + k['posotita']

    # Χτίζουμε τις γραμμές χρονολογικά:
    # Ταξινομούμε αγορές + επιστροφές μαζί κατά ημερομηνία
    # Κάθε επιστροφή εισάγεται αμέσως μετά την αγορά με την οποία χρονολογικά συσχετίζεται

    # Ταξινόμηση αγορών κατά auxon_arithmos (χρονολογική σειρά εισαγωγής)
    agora_list = sorted(agores.values(), key=lambda a: a['auxon'])

    # Για κάθε επιστροφή, βρες την αμέσως προηγούμενη αγορά χρονολογικά
    # που έχει κοινά υλικά
    epi_to_agora = {}  # index επιστροφής → index αγοράς
    for i, e in enumerate(epistrofes):
        best_agora_idx = None
        best_agora_date = None
        for j, agora in enumerate(agora_list):
            common = set(e['ylika'].keys()) & set(agora['ylika'].keys())
            # Η αγορά πρέπει να έχει auxon ΜΙΚΡΟΤΕΡΟ από την επιστροφή
            if common and agora['auxon'] < e['auxon']:
                if best_agora_date is None or agora['auxon'] > (agora_list[best_agora_idx]['auxon'] if best_agora_idx is not None else -1):
                    best_agora_date = agora['imerominia']
                    best_agora_idx = j
        if best_agora_idx is not None:
            epi_to_agora[i] = best_agora_idx

    # Χτίσε τις γραμμές: για κάθε αγορά, βάλε αμέσως μετά τις επιστροφές της
    rows = []
    epi_used = set()
    for j, agora in enumerate(agora_list):
        rows.append(agora)
        for i, e in enumerate(epistrofes):
            if i not in epi_used and epi_to_agora.get(i) == j:
                e['aa'] = 0
                rows.append(e)
                epi_used.add(i)

    # Επιστροφές που δεν συσχετίστηκαν — στο τέλος
    for i, e in enumerate(epistrofes):
        if i not in epi_used:
            e['aa'] = 0
            rows.append(e)

    # Δώσε Α/Α με τη σωστή σειρά τώρα που οι γραμμές είναι στη θέση τους
    for i, row in enumerate(rows, 1):
        row['aa'] = i

    # ── Κατανάλωση ανά parstatiko αγοράς ──────────────────────────────────────
    # kat_by_parst: κλειδί = parstatiko αγοράς
    kat_by_parst = {}

    # Χάρτης ημερομηνία → parstatiko αγοράς (για συσχέτιση ΚΑΤΑΝΑΛΩΣΗ χωρίς parstatiko)
    agora_by_date = {}
    for agora in agora_list:
        agora_by_date[agora['imerominia']] = agora['parstatiko']

    # Καταναλώσεις — συσχέτισε με αγορά βάσει ημερομηνίας αν δεν έχουν parstatiko
    for k in katanaliseis:
        key = k['parstatiko'] if k['parstatiko'] else agora_by_date.get(k['imerominia'], k['imerominia'])
        if key not in kat_by_parst:
            kat_by_parst[key] = {'ylika':{}, 'paratirishis':k['paratirishis'], 'imerominia':k['imerominia'], 'auto':False, 'auxon': k.get('auxon_arithmos', 0)}
        for yid, pos in k['ylika'].items():
            kat_by_parst[key]['ylika'][yid] = kat_by_parst[key]['ylika'].get(yid,0) + pos

    # Για κάθε αγορά χωρίς χειροκίνητη κατανάλωση → αυτόματος υπολογισμός
    # Επιστροφές ανά αγορά — χρήση agora_ref για ακριβή συσχέτιση
    epi_by_agora = {}  # parstatiko αγοράς → {yid: posotita}
    for e in epistrofes:
        # Προτεραιότητα στο agora_ref, fallback στον αλγόριθμο auxon
        agora_parst = e.get('agora_ref') or ''
        if not agora_parst:
            # Fallback: αμέσως προηγούμενη αγορά βάσει auxon
            best = None
            for agora in agora_list:
                common = set(e['ylika'].keys()) & set(agora['ylika'].keys())
                if common and agora['auxon'] < e['auxon']:
                    if best is None or agora['auxon'] > best['auxon']:
                        best = agora
            agora_parst = best['parstatiko'] if best else ''
        if agora_parst:
            if agora_parst not in epi_by_agora:
                epi_by_agora[agora_parst] = {}
            for yid, pos in e['ylika'].items():
                epi_by_agora[agora_parst][yid] = epi_by_agora[agora_parst].get(yid,0) + pos

    for agora in agora_list:
        ap = agora['parstatiko']
        epi = epi_by_agora.get(ap, {})
        if ap not in kat_by_parst:
            # Δεν υπάρχει χειροκίνητη → αυτόματος υπολογισμός: Αγορά − Επιστροφές
            auto_ylika = {}
            for yid, pos in agora['ylika'].items():
                katanalosi = pos - epi.get(yid, 0)
                if katanalosi > 0:
                    auto_ylika[yid] = katanalosi
            if auto_ylika:
                kat_by_parst[ap] = {
                    'ylika': auto_ylika,
                    'paratirishis': 'Αυτόματος υπολογισμός',
                    'imerominia': agora['imerominia'],
                    'auto': True
                }
        else:
            # Υπάρχει χειροκίνητη κατανάλωση
            # Αφαίρεσε επιστροφές ΜΟΝΟ αν η κατανάλωση καταχωρήθηκε ΠΡΙΝ την επιστροφή
            # (δηλαδή η κατανάλωση δεν έχει ήδη αφαιρέσει την επιστροφή)
            if epi:
                # Βρες min auxon κατανάλωσης vs min auxon επιστροφής
                kat_auxon = kat_by_parst[ap].get('auxon', 999999)
                # Συμπεριλάβε ΟΛΕΣ τις επιστροφές που συσχετίστηκαν με αυτή την αγορά
                # (είτε μέσω agora_ref είτε μέσω fallback auxon-αλγόριθμου)
                epi_auxons = [e.get('auxon', 0) for e in epistrofes
                              if (e.get('agora_ref') == ap) or
                                 (not e.get('agora_ref') and
                                  set(e['ylika'].keys()) & set(agora['ylika'].keys()) and
                                  e['auxon'] > agora['auxon'])]
                min_epi_auxon = min(epi_auxons) if epi_auxons else 999999
                # Αφαίρεσε μόνο αν η κατανάλωση έγινε ΠΡΙΝ την επιστροφή
                if kat_auxon < min_epi_auxon:
                    for yid, epi_pos in epi.items():
                        if yid in kat_by_parst[ap]['ylika']:
                            new_pos = kat_by_parst[ap]['ylika'][yid] - epi_pos
                            if new_pos > 0:
                                kat_by_parst[ap]['ylika'][yid] = new_pos
                            else:
                                del kat_by_parst[ap]['ylika'][yid]

    return rows, kat_by_parst


# ─── PDF ─────────────────────────────────────────────────────────────────────

def export_pdf(kiniseis: list, yliko_label: str, period_label: str, font: str = 'iosevka', nonel_mode: str = 'detail') -> bytes:
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4, landscape
    from reportlab.lib.units import cm
    from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, PageBreak
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT

    F, FB = register_fonts(font)
    if not FONT_REGULAR:
        F, FB = 'Helvetica', 'Helvetica-Bold'

    ylika_order = get_ylika_order(kiniseis, nonel_mode)
    rows, kat_by_parst = build_book_rows(kiniseis)

    from database import get_all_ylika as _gay_main
    _all_y = {y['id']: y for y in _gay_main()}
    nonel_delay_map = {}
    if nonel_mode == 'subgroup':
        for k in kiniseis:
            yid = k['yliko_id']
            ydata = _all_y.get(yid, {})
            if not _is_nonel(ydata):
                continue
            gkey = nonel_group_key(ydata)
            vkey = 'NONEL_DLY_{}_{}'.format(*gkey)
            if vkey not in nonel_delay_map:
                nonel_delay_map[vkey] = []
            if yid not in nonel_delay_map[vkey]:
                nonel_delay_map[vkey].append(yid)

    pagesize  = landscape(A4)
    page_w_cm = (pagesize[0] - 1.4*cm) / cm

    TS = ParagraphStyle('t', fontSize=9,  fontName=FB, alignment=TA_CENTER, spaceAfter=2)
    SS = ParagraphStyle('s', fontSize=6,  fontName=F,  alignment=TA_CENTER, spaceAfter=4)
    HS = ParagraphStyle('h', fontSize=5.5,fontName=FB, alignment=TA_CENTER, leading=7, textColor=colors.HexColor('#1C252E'))
    CS = ParagraphStyle('c', fontSize=5.5,fontName=F,  alignment=TA_CENTER, leading=7, textColor=colors.HexColor('#1C252E'))
    RS = ParagraphStyle('r', fontSize=5.5,fontName=F,  alignment=TA_RIGHT,  leading=7, textColor=colors.HexColor('#1C252E'))
    NS = ParagraphStyle('n', fontSize=5.5,fontName=F,  alignment=TA_LEFT,   leading=7, textColor=colors.HexColor('#1C252E'))
    ES = ParagraphStyle('e', fontSize=5.5,fontName=FB, alignment=TA_CENTER, leading=7,
                        textColor=colors.HexColor('#c0392b'))
    ER = ParagraphStyle('er',fontSize=5.5,fontName=FB, alignment=TA_RIGHT,  leading=7,
                        textColor=colors.HexColor('#c0392b'))

    # Χρώματα — ανοιχτό επαγγελματικό γκρι, ίδια και στις 2 σελίδες
    HDR_BG = colors.HexColor('#CDD5DB')  # Τίτλος - λίγο σκουρότερο
    SUB_BG = colors.HexColor('#E4E9ED')
    GRP_BG = colors.HexColor('#DDE3E8')  # ενδιάμεσο tier (group NONEL/subtype) στο ομαδοποιημένο header
    HDR_FG = colors.HexColor('#1C252E')
    ALT_BG = colors.HexColor('#F4F6F8')
    GRID_C = colors.HexColor('#D1D8DE')

    def p(txt, s=None): return Paragraph(str(txt), s or CS)

    MERGED_IDS = {
        'POLADYN':     ([4, 3],  'POLADYN 65X500<br/>POLADYN 38X380', 'Κιλ'),
        'THRYALLIDES': ([5, 10], 'ΒΡΑΔΥΚΑΥΣΤΗ<br/>ΑΚΑΡΙΑΙΑ',          'Μετρ'),
        'KAPSYLIA':    ([9, 33], 'ΚΟΙΝΟΙ ΠΥΡΟΚΡ.<br/>ΗΛΕΚΤΡΙΚΟΙ',     'Τεμ'),
    }
    # Χτίζουμε virtual ylika_order με merged στήλες
    virtual_order = []
    added_merged = set()
    for yid, (on, mo) in ylika_order:
        if yid in {4, 3} and 'POLADYN' not in added_merged:
            virtual_order.append(('POLADYN', ('POLADYN 65X500<br/>POLADYN 38X380', 'Κιλ')))
            added_merged.add('POLADYN')
        elif yid in {5, 10} and 'THRYALLIDES' not in added_merged:
            virtual_order.append(('THRYALLIDES', ('ΒΡΑΔΥΚΑΥΣΤΗ<br/>ΑΚΑΡΙΑΙΑ', 'Μετρ')))
            added_merged.add('THRYALLIDES')
        elif yid in {9, 33} and 'KAPSYLIA' not in added_merged:
            virtual_order.append(('KAPSYLIA', ('ΚΟΙΝΟΙ ΠΥΡΟΚΡ.<br/>ΗΛΕΚΤΡΙΚΟΙ', 'Τεμ')))
            added_merged.add('KAPSYLIA')
        elif yid not in {4, 3, 5, 10, 9, 33}:
            virtual_order.append((yid, (on, mo)))
    ylika_order = virtual_order

    # ── Ομαδοποιημένο header (μόνο σε nonel_mode='detail') ────────────────────
    # Αντί για ένα flat "NONEL SNAPLINE, SL17 (6.0M) (Τεμ)" ανά στήλη (τυλίγεται
    # σε ~7 γραμμές σε τόσο στενές στήλες), 3 επίπεδα: NONEL -> SNAPLINE/UNIDET
    # U500/LP -> SL17/SL42/.../μήκος. Ίδιο μοτίβο και για POLADYN/ΘΡΥΑΛ./ΠΥΡΟΚΡ.
    GROUPED_HEADER = (nonel_mode == 'detail')
    n_fixed = n_nonel = 0
    if GROUPED_HEADER:
        ylika_order, n_fixed, n_nonel, nonel_meta = grouped_header_plan(ylika_order, _all_y)
        PARENT_LEAVES = GROUPED_PARENT_LEAVES
        SHORT_LABEL = {yid: '<br/>'.join(parts) for yid, parts in GROUPED_SHORT_LABEL.items()}

        def build_group_header(LEAD, TAIL, lead_labels, tail_labels):
            n = len(ylika_order)
            row_group = [''] * (LEAD + n + TAIL)
            row_sub   = [''] * (LEAD + n + TAIL)
            row_leaf  = [''] * (LEAD + n + TAIL)
            spans = []
            for i, lbl in enumerate(lead_labels):
                row_group[i] = p(lbl, HS)
                spans.append(('SPAN', (i, 1), (i, 3)))
            for i, lbl in enumerate(tail_labels):
                col = LEAD + n + i
                row_group[col] = p(lbl, HS)
                spans.append(('SPAN', (col, 1), (col, 3)))

            for c in range(n_fixed):
                col = LEAD + c
                yid = ylika_order[c][0]
                if yid in PARENT_LEAVES:
                    parent, (leaf1, leaf2) = PARENT_LEAVES[yid]
                    row_group[col] = p(parent, HS)
                    row_sub[col]   = p(f"{leaf1}<br/>{leaf2}", HS)
                    spans.append(('SPAN', (col, 2), (col, 3)))
                else:
                    on, mo = ylika_order[c][1]
                    row_group[col] = p(SHORT_LABEL.get(yid, on), HS)
                    spans.append(('SPAN', (col, 1), (col, 3)))

            if n_nonel:
                nonel_start = LEAD + n_fixed
                nonel_end   = LEAD + n_fixed + n_nonel - 1
                row_group[nonel_start] = p('NONEL', HS)
                spans.append(('SPAN', (nonel_start, 1), (nonel_end, 1)))

                sub_labels = GROUPED_SUB_LABEL
                run_start = nonel_start
                cur_sub = nonel_meta[ylika_order[n_fixed][0]][0]
                for c in range(nonel_start, nonel_end + 2):
                    yid_c = ylika_order[c - LEAD][0] if c <= nonel_end else None
                    sub = nonel_meta[yid_c][0] if yid_c is not None else None
                    if sub != cur_sub:
                        row_sub[run_start] = p(sub_labels.get(cur_sub, cur_sub), HS)
                        if c - 1 > run_start:
                            spans.append(('SPAN', (run_start, 2), (c - 1, 2)))
                        run_start = c
                        cur_sub = sub

                code_counts = {}
                for c in range(nonel_start, nonel_end + 1):
                    sub, code, length = nonel_meta[ylika_order[c - LEAD][0]]
                    code_counts[(sub, code)] = code_counts.get((sub, code), 0) + 1
                for c in range(nonel_start, nonel_end + 1):
                    sub, code, length = nonel_meta[ylika_order[c - LEAD][0]]
                    if sub == 'UNIDET':
                        row_leaf[c] = p(length, HS)  # ο γονέας λέει ήδη "UNIDET U500"
                    elif code_counts[(sub, code)] > 1:
                        row_leaf[c] = p(f"{code}<br/>{length}", HS)
                    else:
                        row_leaf[c] = p(code, HS)

            return row_group, row_sub, row_leaf, spans

        def grouped_style(spans):
            return [
                ('SPAN', (0, 0), (-1, 0)),
                ('BACKGROUND', (0, 0), (-1, 0), HDR_BG),
                ('BACKGROUND', (0, 1), (-1, 1), SUB_BG),
                ('BACKGROUND', (0, 2), (-1, 2), GRP_BG),
                ('BACKGROUND', (0, 3), (-1, 3), SUB_BG),
                ('TEXTCOLOR', (0, 0), (-1, 0), HDR_FG),
                ('TEXTCOLOR', (0, 1), (-1, 3), colors.black),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                ('GRID', (0, 0), (-1, -1), 0.3, GRID_C),
                ('LINEBELOW', (0, 3), (-1, 3), 1.0, colors.HexColor('#9AAAB5')),
                ('TOPPADDING', (0, 0), (-1, -1), 3),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 3),
                ('LEFTPADDING', (0, 0), (-1, -1), 3),
                ('RIGHTPADDING', (0, 0), (-1, -1), 4),
            ] + spans

        # Μη-ομοιόμορφο πλάτος: πλατύτερες οι στήλες με μεγάλους αριθμούς
        # (Κιλ/Μετρ: AN-FO/EM-EX φτάνουν σε χιλιάδες), στενότερες οι NONEL
        # (Τεμ, σχεδόν πάντα 1-3 ψηφία) — υπολογίζεται μία φορά, ίδιο και
        # στις 2 σελίδες.
        FIXED_W = {1: 1.3, 2: 1.3, 'POLADYN': 0.95, 'THRYALLIDES': 0.85, 'KAPSYLIA': 0.95}
        _L_FIXED_BUDGET = [0.5, 1.7, 1.8, 1.6]
        fixed_total = sum(FIXED_W.get(yid, 1.0) for yid, _ in ylika_order[:n_fixed])
        remaining = page_w_cm - sum(_L_FIXED_BUDGET) - fixed_total
        yw_nonel = max(0.65, remaining / n_nonel) if n_nonel else 0.8
        material_widths = [FIXED_W.get(yid, 1.0) for yid, _ in ylika_order[:n_fixed]] + [yw_nonel] * n_nonel

    n = len(ylika_order)
    if not GROUPED_HEADER:
        yliko_hdrs = [p(f"{on}<br/>({mo})", HS) for _, (on, mo) in ylika_order]

    # ── Σελίδα 1: Αγορές / Επιστροφές ───────────────────────────────────────
    L_FIXED = [0.5, 1.7, 1.8, 1.6]
    if GROUPED_HEADER:
        L_WIDTHS = [L_FIXED[0]*cm, L_FIXED[1]*cm] + [w*cm for w in material_widths] + \
                   [L_FIXED[2]*cm, L_FIXED[3]*cm]
        row_group, row_sub, row_leaf, l_spans = build_group_header(
            2, 2, ['Α/Α', 'Αρ.Άδ./<br/>Εκδ.Αρχή'], ['Ημερ.Αγ./<br/>Αρ.Δελτ.', 'Στοιχεία<br/>Προμηθευτή'])
        l_title = [p('ΒΙΒΛΙΟ ΑΓΟΡΑΣ ΚΑΙ ΚΑΤΑΝΑΛΩΣΗΣ ΕΚΡΗΚΤΙΚΩΝ ΥΛΩΝ — ΑΓΟΡΕΣ / ΕΠΙΣΤΡΟΦΕΣ', HS)] + \
                  [''] * (len(L_WIDTHS)-1)
        l_data  = [l_title, row_group, row_sub, row_leaf]
        l_style = grouped_style(l_spans)
        HDR_ROWS = 4
    else:
        yw_l = max(0.8, (page_w_cm - sum(L_FIXED)) / n) if n else 1.5
        L_WIDTHS = [L_FIXED[0]*cm, L_FIXED[1]*cm] + [yw_l*cm]*n + \
                   [L_FIXED[2]*cm, L_FIXED[3]*cm]
        l_title = [p('ΒΙΒΛΙΟ ΑΓΟΡΑΣ ΚΑΙ ΚΑΤΑΝΑΛΩΣΗΣ ΕΚΡΗΚΤΙΚΩΝ ΥΛΩΝ — ΑΓΟΡΕΣ / ΕΠΙΣΤΡΟΦΕΣ', HS)] + \
                  [''] * (len(L_WIDTHS)-1)
        l_hdr   = [p('Α/Α',HS), p('Αρ.Άδ./<br/>Εκδ.Αρχή',HS)] + yliko_hdrs + \
                  [p('Ημερ.Αγ./<br/>Αρ.Δελτ.',HS), p('Στοιχεία\nΠρομηθευτή',HS)]
        l_data  = [l_title, l_hdr]
        l_style = [
            ('SPAN',         (0,0), (-1,0)),
            ('BACKGROUND',   (0,0), (-1,0), HDR_BG),
            ('BACKGROUND',   (0,1), (-1,1), SUB_BG),
            ('TEXTCOLOR',    (0,0), (-1,0), HDR_FG),        # μαύρο στο ανοιχτό
            ('TEXTCOLOR',    (0,1), (-1,1), colors.black),  # λευκό στο σκούρο
            ('ALIGN',        (0,0), (-1,-1), 'CENTER'),
            ('VALIGN',       (0,0), (-1,-1), 'MIDDLE'),
            ('GRID',         (0,0), (-1,-1), 0.3, GRID_C),
            ('LINEBELOW',    (0,1), (-1,1), 1.0, colors.HexColor('#9AAAB5')),
            ('TOPPADDING',   (0,0), (-1,-1), 3),
            ('BOTTOMPADDING',(0,0), (-1,-1), 3),
            ('LEFTPADDING',  (0,0), (-1,-1), 3),
            ('RIGHTPADDING', (0,0), (-1,-1), 4),
        ]
        HDR_ROWS = 2

    n_ylika_cols = len(ylika_order)
    num_l_start = 2  # col index αρχή αριθμών αριστερά
    num_l_end   = 2 + n_ylika_cols - 1
    l_style.append(('ALIGN', (num_l_start, HDR_ROWS), (num_l_end, -1), 'RIGHT'))

    for ri, row in enumerate(rows, HDR_ROWS):
        is_epi = row['type'] == 'epistrofi'
        txt_s  = ES if is_epi else CS
        num_s  = ER if is_epi else RS
        adeia_ekd = row['adeia']
        if row.get('ekdousa'):
            adeia_ekd += f"<br/>{row['ekdousa']}"

        cells = [
            p(str(row['aa']) if row['aa'] else '', txt_s),
            p(adeia_ekd, txt_s),
        ]
        _nonel_sums = get_nonel_sums(row['ylika'], _all_y, nonel_mode) if nonel_mode not in ('detail', 'subgroup') else {}
        _MERGED = {'POLADYN': [4, 3], 'THRYALLIDES': [5, 10], 'KAPSYLIA': [9, 33]}
        for yid, _ in ylika_order:
            if isinstance(yid, str) and yid.startswith('NONEL_DLY_'):
                ids = nonel_delay_map.get(yid, [])
                total = sum(row['ylika'].get(i, 0) for i in ids) or None
                cells.append(p(fmt_num(total) if total else '—', num_s))
            elif isinstance(yid, str) and yid.startswith('NONEL_'):
                if yid == 'NONEL_TOTAL':
                    v = _nonel_sums.get('TOTAL', 0) or None
                else:
                    sub = yid.replace('NONEL_', '')
                    v = _nonel_sums.get(sub, 0) or None
                cells.append(p(fmt_num(v) if v else '—', num_s))
            elif isinstance(yid, str) and yid in _MERGED:
                ids = _MERGED[yid]
                parts = [fmt_num(row['ylika'].get(i)) if row['ylika'].get(i) else '—' for i in ids]
                cells.append(p('<br/>'.join(parts), num_s))
            else:
                v = row['ylika'].get(yid)
                cells.append(p(fmt_num(v) if v else '—', num_s))
        cells += [
            p(f"{fmt_date(row['imerominia'])}<br/>{row['parstatiko']}", txt_s),
            p(row['promitheftis'], txt_s),
        ]
        l_data.append(cells)

        if ri % 2 == 0:
            l_style.append(('BACKGROUND', (0,ri), (-1,ri), ALT_BG))

    # Σταθερό ύψος γραμμών — ίδιο και στους 2 πίνακες για αντικρυστή εμφάνιση
    ROW_H = 22  # points — αρκετό για 2 γραμμές κειμένου
    if GROUPED_HEADER:
        header_heights = [16, 12, 12, 20]
    else:
        header_heights = [28, 28]
    n_data_rows = len(l_data) - HDR_ROWS
    row_heights = header_heights + [ROW_H] * n_data_rows

    left_t = Table(l_data, colWidths=L_WIDTHS, repeatRows=HDR_ROWS, rowHeights=row_heights)
    left_t.setStyle(TableStyle(l_style))

    # ── Σελίδα 2: Καταναλώσεις ───────────────────────────────────────────────
    R_FIXED = [1.5, 1.5, 1.8]
    if GROUPED_HEADER:
        R_WIDTHS = [R_FIXED[0]*cm, R_FIXED[1]*cm] + [w*cm for w in material_widths] + [R_FIXED[2]*cm]
        row_group2, row_sub2, row_leaf2, r_spans = build_group_header(
            2, 1, ['Ημερ.<br/>Εισαγ.<br/>Αποθ.', 'Ημερ.<br/>Κατανάλ.'], ['Παρατηρήσεις'])
        r_title = [p('ΒΙΒΛΙΟ ΑΓΟΡΑΣ ΚΑΙ ΚΑΤΑΝΑΛΩΣΗΣ ΕΚΡΗΚΤΙΚΩΝ ΥΛΩΝ — ΚΑΤΑΝΑΛΩΣΕΙΣ', HS)] + \
                  [''] * (len(R_WIDTHS)-1)
        r_data  = [r_title, row_group2, row_sub2, row_leaf2]
        r_style = grouped_style(r_spans)
    else:
        yw_r = max(0.8, (page_w_cm - sum(R_FIXED)) / n) if n else 1.5
        R_WIDTHS = [R_FIXED[0]*cm, R_FIXED[1]*cm] + [yw_r*cm]*n + [R_FIXED[2]*cm]
        r_title = [p('ΒΙΒΛΙΟ ΑΓΟΡΑΣ ΚΑΙ ΚΑΤΑΝΑΛΩΣΗΣ ΕΚΡΗΚΤΙΚΩΝ ΥΛΩΝ — ΚΑΤΑΝΑΛΩΣΕΙΣ', HS)] + \
                  [''] * (len(R_WIDTHS)-1)
        r_hdr   = [p('Ημερ.\nΕισαγ.\nΑποθ.',HS), p('Ημερ.\nΚατανάλ.',HS)] + \
                  yliko_hdrs + [p('Παρατηρήσεις',HS)]
        r_data  = [r_title, r_hdr]
        r_style = [
            ('SPAN',         (0,0), (-1,0)),
            ('BACKGROUND',   (0,0), (-1,0), HDR_BG),
            ('BACKGROUND',   (0,1), (-1,1), SUB_BG),
            ('TEXTCOLOR',    (0,0), (-1,0), HDR_FG),
            ('TEXTCOLOR',    (0,1), (-1,1), colors.black),
            ('ALIGN',        (0,0), (-1,-1), 'CENTER'),
            ('VALIGN',       (0,0), (-1,-1), 'MIDDLE'),
            ('GRID',         (0,0), (-1,-1), 0.3, GRID_C),
            ('LINEBELOW',    (0,1), (-1,1), 1.0, colors.HexColor('#9AAAB5')),
            ('TOPPADDING',   (0,0), (-1,-1), 3),
            ('BOTTOMPADDING',(0,0), (-1,-1), 3),
            ('LEFTPADDING',  (0,0), (-1,-1), 3),
            ('RIGHTPADDING', (0,0), (-1,-1), 4),
        ]
    r_style.append(('ALIGN', (2, HDR_ROWS), (1+n_ylika_cols, -1), 'RIGHT'))

    kat_used = set()
    for ri, row in enumerate(rows, HDR_ROWS):
        is_epi = row['type'] == 'epistrofi'
        ap     = row.get('parstatiko', '')
        imer   = row['imerominia']

        if not is_epi and ap not in kat_used:
            kat = kat_by_parst.get(ap, kat_by_parst.get(imer, {}))
            kat_used.add(ap)
            kat_imer = kat.get('imerominia', imer)
            cells = [p(fmt_date(kat_imer), CS), p(fmt_date(kat_imer), CS)]
            _nonel_sums2 = get_nonel_sums(kat.get('ylika',{}), _all_y, nonel_mode) if nonel_mode not in ('detail', 'subgroup') else {}
            _MERGED = {'POLADYN': [4, 3], 'THRYALLIDES': [5, 10], 'KAPSYLIA': [9, 33]}
            kat_ylika = kat.get('ylika', {})
            for yid, _ in ylika_order:
                if isinstance(yid, str) and yid.startswith('NONEL_DLY_'):
                    ids = nonel_delay_map.get(yid, [])
                    total = sum(kat_ylika.get(i, 0) for i in ids) or None
                    cells.append(p(fmt_num(total) if total else '—', RS))
                elif isinstance(yid, str) and yid.startswith('NONEL_'):
                    if yid == 'NONEL_TOTAL':
                        v = _nonel_sums2.get('TOTAL', 0) or None
                    else:
                        sub = yid.replace('NONEL_', '')
                        v = _nonel_sums2.get(sub, 0) or None
                    cells.append(p(fmt_num(v) if v else '—', RS))
                elif isinstance(yid, str) and yid in _MERGED:
                    ids = _MERGED[yid]
                    parts = [fmt_num(kat_ylika.get(i)) if kat_ylika.get(i) else '—' for i in ids]
                    cells.append(p('<br/>'.join(parts), RS))
                else:
                    v = kat_ylika.get(yid)
                    cells.append(p(fmt_num(v) if v else '—', RS))
            cells.append(p(kat.get('paratirishis',''), CS))
            if ri % 2 == 0:
                r_style.append(('BACKGROUND', (0,ri), (-1,ri), ALT_BG))
        else:
            cells = [p('',CS), p('',CS)] + [p('',CS)]*n + [p('',CS)]

        r_data.append(cells)

    right_t = Table(r_data, colWidths=R_WIDTHS, repeatRows=HDR_ROWS, rowHeights=row_heights)
    right_t.setStyle(TableStyle(r_style))

    # ── Build ─────────────────────────────────────────────────────────────────
    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=pagesize,
                            leftMargin=0.7*cm, rightMargin=0.7*cm,
                            topMargin=1.2*cm, bottomMargin=1.2*cm)
    doc.build([left_t, PageBreak(), right_t])
    return buf.getvalue()


# ─── Excel ────────────────────────────────────────────────────────────────────

def export_excel(kiniseis: list, yliko_label: str, period_label: str, nonel_mode: str = 'detail') -> bytes:
    import openpyxl
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    from openpyxl.utils import get_column_letter

    rows, kat_by_parst = build_book_rows(kiniseis)

    from database import get_all_ylika as _gay_xl
    _all_y = {y['id']: y for y in _gay_xl()}

    _MERGED_IDS = {'POLADYN': [4, 3], 'THRYALLIDES': [5, 10], 'KAPSYLIA': [9, 33]}
    _MERGED_HDR = {
        'POLADYN':     ('POLADYN 65X500 / POLADYN 38X380', 'Κιλ'),
        'THRYALLIDES': ('ΒΡΑΔΥΚΑΥΣΤΗ / ΑΚΑΡΙΑΙΑ',          'Μετρ'),
        'KAPSYLIA':    ('ΚΟΙΝΟΙ ΠΥΡΟΚΡ. / ΗΛΕΚΤΡΙΚΟΙ',    'Τεμ'),
    }

    def build_virtual_order(mode):
        raw = get_ylika_order(kiniseis, mode)
        vo, added = [], set()
        for yid, (on, mo) in raw:
            if yid in {4, 3} and 'POLADYN' not in added:
                vo.append(('POLADYN', _MERGED_HDR['POLADYN'])); added.add('POLADYN')
            elif yid in {5, 10} and 'THRYALLIDES' not in added:
                vo.append(('THRYALLIDES', _MERGED_HDR['THRYALLIDES'])); added.add('THRYALLIDES')
            elif yid in {9, 33} and 'KAPSYLIA' not in added:
                vo.append(('KAPSYLIA', _MERGED_HDR['KAPSYLIA'])); added.add('KAPSYLIA')
            elif yid not in {4, 3, 5, 10, 9, 33}:
                vo.append((yid, (on, mo)))
        return vo

    def build_nonel_delay_map(mode, vo):
        if mode != 'subgroup':
            return {}
        dm = {}
        for k in kiniseis:
            yid = k['yliko_id']
            ydata = _all_y.get(yid, {})
            if not _is_nonel(ydata):
                continue
            gkey = nonel_group_key(ydata)
            vkey = 'NONEL_DLY_{}_{}'.format(*gkey)
            if vkey not in dm:
                dm[vkey] = []
            if yid not in dm[vkey]:
                dm[vkey].append(yid)
        return dm

    def yval(yid, ylika_dict, delay_map):
        if isinstance(yid, str) and yid.startswith('NONEL_DLY_'):
            ids = delay_map.get(yid, [])
            v = sum(ylika_dict.get(i, 0) for i in ids) or None
            return v, False
        if isinstance(yid, str) and yid == 'NONEL_TOTAL':
            v = get_nonel_sums(ylika_dict, _all_y, 'total').get('TOTAL') or None
            return v, False
        if isinstance(yid, str) and yid.startswith('NONEL_'):
            sub = yid.replace('NONEL_', '')
            v = get_nonel_sums(ylika_dict, _all_y, 'grouped').get(sub) or None
            return v, False
        if isinstance(yid, str) and yid in _MERGED_IDS:
            parts = [ylika_dict.get(i) for i in _MERGED_IDS[yid]]
            return '\n'.join(fmt_num(v) if v else '—' for v in parts), True
        return ylika_dict.get(yid), False

    wb = openpyxl.Workbook()
    wb.remove(wb.active)  # αφαιρούμε το default φύλλο

    thin   = Side(style='thin', color="AABBD0")
    border = Border(left=thin, right=thin, top=thin, bottom=thin)

    MODES = [
        ('detail',   'Αναλυτικό'),
        ('subgroup', 'Ανά καθυστέρηση'),
        ('grouped',  'Ανά τύπο NONEL'),
        ('total',    'NONEL Σύνολο'),
    ]

    for mode, sheet_label in MODES:
        vo       = build_virtual_order(mode)
        grouped  = (mode == 'detail')
        n_fixed = n_nonel = 0
        nonel_meta = {}
        if grouped:
            vo, n_fixed, n_nonel, nonel_meta = grouped_header_plan(vo, _all_y)
        dly_map  = build_nonel_delay_map(mode, vo)
        n        = len(vo)
        total_cols = max(2 + n + 2, 2 + n + 1)  # αγορές έχει +2, καταν. +1
        HDR_ROWS = 4 if grouped else 2  # τελευταία γραμμή header (η γραμμή τίτλου είναι πάντα row 1)

        ws = wb.create_sheet(sheet_label)

        hdr_font  = Font(name='Iosevka', bold=True, color="FFFFFF", size=9)
        data_font = Font(name='Iosevka', size=9)
        alt_fill  = PatternFill("solid", fgColor="EEF2F7")
        red_fill  = PatternFill("solid", fgColor="FFF0F0")
        red_font  = Font(name='Iosevka', color="C53030", bold=True, size=9)
        navy_fill = PatternFill("solid", fgColor="1A365D")
        lnav_fill = PatternFill("solid", fgColor="4A7FC1")
        lnav_fill2= PatternFill("solid", fgColor="7FA5D6")
        grn_fill  = PatternFill("solid", fgColor="2D6A4F")
        lgrn_fill = PatternFill("solid", fgColor="52B788")
        lgrn_fill2= PatternFill("solid", fgColor="8FD3AC")
        alt2_fill = PatternFill("solid", fgColor="D8F3DC")

        # Layout: Αγορές αριστερά | κενή στήλη | Καταναλώσεις δεξιά
        L_COLS  = 2 + n + 2   # Α/Α, Άδεια, υλικά, Ημερ/Δελτ, Προμηθ
        GAP     = 1
        R_START = L_COLS + GAP + 1   # 1-indexed αρχή δεξιού πίνακα
        R_COLS  = 2 + n + 1   # Ημερ×2, υλικά, Παρατηρήσεις

        # ── Τίτλοι row 1 ─────────────────────────────────────────────────────
        ws.merge_cells(f'A1:{get_column_letter(L_COLS)}1')
        ws['A1'] = "ΒΙΒΛΙΟ ΑΓΟΡΑΣ ΚΑΙ ΚΑΤΑΝΑΛΩΣΗΣ ΕΚΡΗΚΤΙΚΩΝ ΥΛΩΝ — ΑΓΟΡΕΣ / ΕΠΙΣΤΡΟΦΕΣ"
        ws['A1'].fill = navy_fill
        ws['A1'].font = Font(name='Iosevka', bold=True, color="FFFFFF", size=11)
        ws['A1'].alignment = Alignment(horizontal='center')

        rc1 = get_column_letter(R_START)
        rc2 = get_column_letter(R_START + R_COLS - 1)
        ws.merge_cells(f'{rc1}1:{rc2}1')
        ws.cell(row=1, column=R_START).value = \
            "ΒΙΒΛΙΟ ΑΓΟΡΑΣ ΚΑΙ ΚΑΤΑΝΑΛΩΣΗΣ ΕΚΡΗΚΤΙΚΩΝ ΥΛΩΝ — ΚΑΤΑΝΑΛΩΣΕΙΣ"
        ws.cell(row=1, column=R_START).fill = grn_fill
        ws.cell(row=1, column=R_START).font = Font(name='Iosevka', bold=True, color="FFFFFF", size=11)
        ws.cell(row=1, column=R_START).alignment = Alignment(horizontal='center')

        def set_cell(row, col, value, fill, wrap=True):
            cell = ws.cell(row=row, column=col, value=value)
            cell.fill = fill; cell.font = hdr_font
            cell.alignment = Alignment(horizontal='center', vertical='center', wrap_text=wrap)
            cell.border = border
            return cell

        def build_group_header_xl(start_col, lead_labels, tail_labels, fill_a, fill_b):
            """3-επίπεδο header (γονέας/subtype/leaf) — μόνο για nonel_mode='detail'.
            Ίδιο μοτίβο με το ομαδοποιημένο header του PDF export_pdf."""
            col = start_col
            for lbl in lead_labels:
                set_cell(2, col, lbl, fill_a)
                ws.merge_cells(start_row=2, start_column=col, end_row=4, end_column=col)
                col += 1
            for c in range(n_fixed):
                yid, (on, mo) = vo[c]
                if yid in GROUPED_PARENT_LEAVES:
                    parent, (leaf1, leaf2) = GROUPED_PARENT_LEAVES[yid]
                    set_cell(2, col, parent, fill_a)
                    set_cell(3, col, f"{leaf1}\n{leaf2}", fill_b)
                    ws.merge_cells(start_row=3, start_column=col, end_row=4, end_column=col)
                else:
                    label = '\n'.join(GROUPED_SHORT_LABEL.get(yid, (on,)))
                    set_cell(2, col, label, fill_a)
                    ws.merge_cells(start_row=2, start_column=col, end_row=4, end_column=col)
                col += 1
            if n_nonel:
                nonel_col0 = col
                set_cell(2, nonel_col0, 'NONEL', fill_a)
                ws.merge_cells(start_row=2, start_column=nonel_col0, end_row=2, end_column=nonel_col0 + n_nonel - 1)

                run_start = nonel_col0
                cur_sub = nonel_meta[vo[n_fixed][0]][0]
                for i in range(n_fixed, n_fixed + n_nonel + 1):
                    cur_col = nonel_col0 + (i - n_fixed)
                    yid_i = vo[i][0] if i < n_fixed + n_nonel else None
                    sub = nonel_meta[yid_i][0] if yid_i is not None else None
                    if sub != cur_sub:
                        set_cell(3, run_start, GROUPED_SUB_LABEL.get(cur_sub, cur_sub), fill_b)
                        if cur_col - 1 > run_start:
                            ws.merge_cells(start_row=3, start_column=run_start, end_row=3, end_column=cur_col - 1)
                        run_start = cur_col
                        cur_sub = sub

                code_counts = {}
                for i in range(n_fixed, n_fixed + n_nonel):
                    sub, code, length = nonel_meta[vo[i][0]]
                    code_counts[(sub, code)] = code_counts.get((sub, code), 0) + 1
                for i in range(n_fixed, n_fixed + n_nonel):
                    sub, code, length = nonel_meta[vo[i][0]]
                    col_i = nonel_col0 + (i - n_fixed)
                    if sub == 'UNIDET':
                        leaf_txt = length
                    elif code_counts[(sub, code)] > 1:
                        leaf_txt = f"{code}\n{length}"
                    else:
                        leaf_txt = code
                    set_cell(4, col_i, leaf_txt, fill_b)
                col = nonel_col0 + n_nonel
            for lbl in tail_labels:
                set_cell(2, col, lbl, fill_a)
                ws.merge_cells(start_row=2, start_column=col, end_row=4, end_column=col)
                col += 1

        # ── Headers ───────────────────────────────────────────────────────────
        if grouped:
            build_group_header_xl(1, ['Α/Α', 'Αρ.Άδ./\nΕκδ.Αρχή'],
                                   ['Ημερ.Αγ./\nΑρ.Δελτ.', 'Στοιχεία\nΠρομηθευτή'],
                                   lnav_fill, lnav_fill2)
            build_group_header_xl(R_START, ['Ημερ.\nΕισαγ.\nΑποθ.', 'Ημερ.\nΚατανάλ.'],
                                   ['Παρατηρήσεις'], lgrn_fill, lgrn_fill2)

            widths_l = [5, 14] + [12 if vo[c][0] in GROUPED_PARENT_LEAVES or vo[c][0] in GROUPED_SHORT_LABEL
                                   else 8 for c in range(n)] + [12, 16]
            for ci, w in enumerate(widths_l, 1):
                ws.column_dimensions[get_column_letter(ci)].width = w
            widths_r = [12, 12] + widths_l[2:2+n] + [20]
            for ci, w in enumerate(widths_r, R_START):
                ws.column_dimensions[get_column_letter(ci)].width = w

            for r in (2, 3, 4):
                ws.row_dimensions[r].height = 16
        else:
            hdrs_l   = ['Α/Α', 'Αρ.Άδ./Εκδ.Αρχή'] + \
                       [f"{on}\n({mo})" for _, (on, mo) in vo] + \
                       ['Ημερ.Αγ./\nΑρ.Δελτ.', 'Στοιχεία\nΠρομηθευτή']
            widths_l = [5, 18] + [13]*n + [12, 18]

            for ci, (h, w) in enumerate(zip(hdrs_l, widths_l), 1):
                cell = ws.cell(row=2, column=ci, value=h)
                cell.fill = lnav_fill; cell.font = hdr_font
                cell.alignment = Alignment(horizontal='center', wrap_text=True)
                cell.border = border
                ws.column_dimensions[get_column_letter(ci)].width = w

            hdrs_r  = ['Ημερ.\nΕισαγ.\nΑποθ.', 'Ημερ.\nΚατανάλ.'] + \
                      [f"{on}\n({mo})" for _, (on, mo) in vo] + ['Παρατηρήσεις']
            widths_r = [14, 14] + [13]*n + [20]

            for ci, (h, w) in enumerate(zip(hdrs_r, widths_r), R_START):
                cell = ws.cell(row=2, column=ci, value=h)
                cell.fill = lgrn_fill; cell.font = hdr_font
                cell.alignment = Alignment(horizontal='center', wrap_text=True)
                cell.border = border
                ws.column_dimensions[get_column_letter(ci)].width = w

            ws.row_dimensions[2].height = 32

        # ── Data rows (ίδιο ri και για τους δύο πίνακες) ─────────────────────
        kat_used = set()
        for ri, row in enumerate(rows, HDR_ROWS + 1):
            is_epi = row['type'] == 'epistrofi'
            fill_l = red_fill if is_epi else (alt_fill if ri % 2 == 0 else None)
            font_l = red_font if is_epi else None
            adeia  = row['adeia'] + (f"\n{row['ekdousa']}" if row.get('ekdousa') else '')

            # — Αριστερός πίνακας (Αγορές) —
            ci = 1
            for v in [row['aa'] or '', adeia]:
                cell = ws.cell(row=ri, column=ci, value=v)
                cell.border = border
                if fill_l: cell.fill = fill_l
                cell.font = font_l or data_font
                cell.alignment = Alignment(horizontal='center', wrap_text=True)
                ci += 1
            for yid, _ in vo:
                v, is_text = yval(yid, row['ylika'], dly_map)
                cell = ws.cell(row=ri, column=ci, value=v)
                cell.border = border
                if fill_l: cell.fill = fill_l
                cell.font = font_l or data_font
                if is_text or not isinstance(v, (int, float)):
                    cell.alignment = Alignment(horizontal='center', wrap_text=True)
                else:
                    cell.number_format = '#,##0.000'
                    cell.alignment = Alignment(horizontal='right')
                ci += 1
            for v in [f"{fmt_date(row['imerominia'])}\n{row['parstatiko']}", row['promitheftis']]:
                cell = ws.cell(row=ri, column=ci, value=v)
                cell.border = border
                if fill_l: cell.fill = fill_l
                cell.font = font_l or data_font
                cell.alignment = Alignment(horizontal='center', wrap_text=True)
                ci += 1

            # — Δεξιός πίνακας (Καταναλώσεις) —
            ap   = row.get('parstatiko', '')
            imer = row['imerominia']
            if not is_epi and ap not in kat_used:
                kat = kat_by_parst.get(ap, kat_by_parst.get(imer, {}))
                kat_used.add(ap)
                kat_imer  = kat.get('imerominia', imer)
                kat_ylika = kat.get('ylika', {})
                fill_r = alt2_fill if ri % 2 == 0 else None

                ci = R_START
                for v in [fmt_date(kat_imer), fmt_date(kat_imer)]:
                    cell = ws.cell(row=ri, column=ci, value=v)
                    cell.border = border
                    if fill_r: cell.fill = fill_r
                    cell.font = data_font
                    cell.alignment = Alignment(horizontal='center')
                    ci += 1
                for yid, _ in vo:
                    v, is_text = yval(yid, kat_ylika, dly_map)
                    cell = ws.cell(row=ri, column=ci, value=v)
                    cell.border = border
                    if fill_r: cell.fill = fill_r
                    cell.font = data_font
                    if is_text or not isinstance(v, (int, float)):
                        cell.alignment = Alignment(horizontal='center', wrap_text=True)
                    else:
                        cell.number_format = '#,##0.000'
                        cell.alignment = Alignment(horizontal='right')
                    ci += 1
                cell = ws.cell(row=ri, column=ci, value=kat.get('paratirishis', ''))
                cell.border = border
                if fill_r: cell.fill = fill_r
                cell.font = data_font
                cell.alignment = Alignment(horizontal='center')

    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


# ─── Word (docx) ─────────────────────────────────────────────────────────────

def export_docx(kiniseis: list, yliko_label: str, period_label: str, nonel_mode: str = 'detail') -> bytes:
    from docx import Document
    from docx.shared import Cm, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    rows, kat_by_parst = build_book_rows(kiniseis)

    from database import get_all_ylika as _gay
    _all_y = {y['id']: y for y in _gay()}

    # Build virtual_order με merged στήλες (ίδια λογική PDF/Excel)
    raw_order = get_ylika_order(kiniseis, nonel_mode)
    virtual_order = []
    added_merged = set()
    for yid, (on, mo) in raw_order:
        if yid in {4, 3} and 'POLADYN' not in added_merged:
            virtual_order.append(('POLADYN', ('POLADYN 65X500 / POLADYN 38X380', 'Κιλ')))
            added_merged.add('POLADYN')
        elif yid in {5, 10} and 'THRYALLIDES' not in added_merged:
            virtual_order.append(('THRYALLIDES', ('ΒΡΑΔΥΚΑΥΣΤΗ / ΑΚΑΡΙΑΙΑ', 'Μετρ')))
            added_merged.add('THRYALLIDES')
        elif yid in {9, 33} and 'KAPSYLIA' not in added_merged:
            virtual_order.append(('KAPSYLIA', ('ΚΟΙΝΟΙ ΠΥΡΟΚΡ. / ΗΛΕΚΤΡΙΚΟΙ', 'Τεμ')))
            added_merged.add('KAPSYLIA')
        elif yid not in {4, 3, 5, 10, 9, 33}:
            virtual_order.append((yid, (on, mo)))

    def is_nonel(yid):
        if isinstance(yid, str):
            return 'NONEL' in yid
        return _is_nonel(_all_y.get(yid, {}))

    # ── Ομαδοποιημένο header (μόνο σε nonel_mode='detail') — ίδιο μοτίβο με
    # PDF/Excel: ένας ενιαίος πίνακας (όχι πλέον χωρισμένος σε "σταθερά υλικά"
    # + ξεχωριστό "NONEL" με διπλότυπες στήλες αναφοράς).
    columns = virtual_order
    grouped = (nonel_mode == 'detail')
    n_fixed = n_nonel = 0
    nonel_meta = {}
    if grouped:
        columns, n_fixed, n_nonel, nonel_meta = grouped_header_plan(columns, _all_y)

    non_nonel  = [(yid, lbl) for yid, lbl in columns if not is_nonel(yid)]
    nonel_cols = [(yid, lbl) for yid, lbl in columns if is_nonel(yid)]

    nonel_delay_map = {}
    if nonel_mode == 'subgroup':
        for k in kiniseis:
            yid = k['yliko_id']
            ydata = _all_y.get(yid, {})
            if not _is_nonel(ydata):
                continue
            gkey = nonel_group_key(ydata)
            vkey = 'NONEL_DLY_{}_{}'.format(*gkey)
            if vkey not in nonel_delay_map:
                nonel_delay_map[vkey] = []
            if yid not in nonel_delay_map[vkey]:
                nonel_delay_map[vkey].append(yid)

    _MRG = {'POLADYN': [4, 3], 'THRYALLIDES': [5, 10], 'KAPSYLIA': [9, 33]}

    def get_val(yid, ylika_dict):
        if isinstance(yid, str) and yid.startswith('NONEL_DLY_'):
            ids = nonel_delay_map.get(yid, [])
            v = sum(ylika_dict.get(i, 0) for i in ids)
            return fmt_num(v) if v else '—'
        if isinstance(yid, str) and yid == 'NONEL_TOTAL':
            v = get_nonel_sums(ylika_dict, _all_y, 'total').get('TOTAL', 0)
            return fmt_num(v) if v else '—'
        if isinstance(yid, str) and yid.startswith('NONEL_'):
            sub = yid.replace('NONEL_', '')
            v = get_nonel_sums(ylika_dict, _all_y, 'grouped').get(sub, 0)
            return fmt_num(v) if v else '—'
        if isinstance(yid, str) and yid in _MRG:
            parts = [fmt_num(ylika_dict.get(i)) if ylika_dict.get(i) else '—' for i in _MRG[yid]]
            return '\n'.join(parts)
        v = ylika_dict.get(yid)
        return fmt_num(v) if v else '—'

    def hex_rgb(h):
        h = h.lstrip('#')
        return tuple(int(h[i:i+2], 16) for i in (0, 2, 4))

    def _shd(cell, fill):
        tcPr = cell._tc.get_or_add_tcPr()
        for x in tcPr.findall(qn('w:shd')):
            tcPr.remove(x)
        el = OxmlElement('w:shd')
        el.set(qn('w:val'), 'clear')
        el.set(qn('w:color'), 'auto')
        el.set(qn('w:fill'), fill)
        tcPr.append(el)

    def _write(cell, text, bold=False, size=7, fg=None, align='center'):
        para = cell.paragraphs[0]
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(0)
        para.alignment = {'center': WD_ALIGN_PARAGRAPH.CENTER,
                          'right':  WD_ALIGN_PARAGRAPH.RIGHT,
                          'left':   WD_ALIGN_PARAGRAPH.LEFT}.get(align, WD_ALIGN_PARAGRAPH.CENTER)
        p_elem = para._p
        for r_elem in list(p_elem.findall(qn('w:r'))):
            p_elem.remove(r_elem)
        rgb = hex_rgb(fg) if fg else None
        for i, line in enumerate(str(text).split('\n')):
            if i > 0:
                br_r = OxmlElement('w:r')
                br_r.append(OxmlElement('w:br'))
                p_elem.append(br_r)
            run = para.add_run(line)
            run.font.name = 'Iosevka'
            run.font.size = Pt(size)
            run.font.bold = bold
            if rgb:
                run.font.color.rgb = RGBColor(*rgb)

    def _tcW(cell, w_cm):
        tcPr = cell._tc.get_or_add_tcPr()
        for x in list(tcPr.findall(qn('w:tcW'))):
            tcPr.remove(x)
        el = OxmlElement('w:tcW')
        el.set(qn('w:w'), str(int(w_cm * 567)))
        el.set(qn('w:type'), 'dxa')
        tcPr.append(el)

    def _vmid(cell):
        tcPr = cell._tc.get_or_add_tcPr()
        el = OxmlElement('w:vAlign')
        el.set(qn('w:val'), 'center')
        tcPr.append(el)

    def _setup_grid(table, widths_cm):
        tbl = table._tbl
        for x in list(tbl.findall(qn('w:tblGrid'))):
            tbl.remove(x)
        tblGrid = OxmlElement('w:tblGrid')
        for w in widths_cm:
            gc = OxmlElement('w:gridCol')
            gc.set(qn('w:w'), str(int(w * 567)))
            tblGrid.append(gc)
        tblPr = tbl.find(qn('w:tblPr'))
        if tblPr is None:
            tblPr = OxmlElement('w:tblPr')
            tbl.insert(0, tblPr)
        tblPr.addnext(tblGrid)
        lay = OxmlElement('w:tblLayout')
        lay.set(qn('w:type'), 'fixed')
        tblPr.append(lay)

    HDR_BG = 'CDD5DB'
    SUB_BG = 'E4E9ED'
    GRP_BG = 'DDE3E8'
    EPY_BG = 'FFF0F0'
    EPY_FG = 'C0392B'
    ALT_BG = 'F4F6F8'

    doc = Document()
    sect = doc.sections[0]
    sect.page_width    = Cm(29.7)
    sect.page_height   = Cm(21.0)
    sect.left_margin   = Cm(1.0)
    sect.right_margin  = Cm(1.0)
    sect.top_margin    = Cm(1.5)
    sect.bottom_margin = Cm(1.5)
    pgSz = sect._sectPr.find(qn('w:pgSz'))
    if pgSz is not None:
        pgSz.set(qn('w:orient'), 'landscape')

    PAGE_W = 27.7  # usable cm

    def add_title(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(3)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        r.font.name = 'Iosevka'
        r.font.bold = True
        r.font.size = Pt(10)

    def material_widths_for(cols):
        """Πλάτη ανά υλικό — πλατύτερα τα Κιλ/Μετρ με μεγάλους αριθμούς
        (AN-FO/EM-EX), στενότερα τα NONEL Τεμ (nonel_mode='detail')."""
        if not grouped:
            n = len(cols)
            return [max(1.5, (PAGE_W - 8.8) / max(n, 1))] * n
        FIXED_W = {1: 1.7, 2: 1.7, 'POLADYN': 1.4, 'THRYALLIDES': 1.2, 'KAPSYLIA': 1.4}
        fixed_total = sum(FIXED_W.get(yid, 1.3) for yid, _ in cols[:n_fixed])
        remaining = PAGE_W - 8.8 - fixed_total
        yw_nonel = max(0.9, remaining / n_nonel) if n_nonel else 1.0
        return [FIXED_W.get(yid, 1.3) for yid, _ in cols[:n_fixed]] + [yw_nonel] * n_nonel

    def build_group_header(tbl, columns_subset, lead_labels, tail_labels):
        """3 γραμμές header (γονέας/subtype/leaf) στις γραμμές 0,1,2 του
        πίνακα — ίδιο μοτίβο με το ομαδοποιημένο header PDF/Excel."""
        col = 0
        for lbl in lead_labels:
            m = tbl.cell(0, col).merge(tbl.cell(2, col))
            _write(m, lbl, bold=True, size=7); _shd(m, HDR_BG); _vmid(m)
            col += 1
        for c in range(n_fixed):
            yid, (on, mo) = columns_subset[c]
            if yid in GROUPED_PARENT_LEAVES:
                parent, (leaf1, leaf2) = GROUPED_PARENT_LEAVES[yid]
                cp = tbl.cell(0, col)
                _write(cp, parent, bold=True, size=7); _shd(cp, HDR_BG); _vmid(cp)
                m = tbl.cell(1, col).merge(tbl.cell(2, col))
                _write(m, f"{leaf1}\n{leaf2}", bold=True, size=7); _shd(m, SUB_BG); _vmid(m)
            else:
                label = '\n'.join(GROUPED_SHORT_LABEL.get(yid, (on,)))
                m = tbl.cell(0, col).merge(tbl.cell(2, col))
                _write(m, label, bold=True, size=7); _shd(m, HDR_BG); _vmid(m)
            col += 1
        if n_nonel:
            nonel_col0 = col
            m = tbl.cell(0, nonel_col0).merge(tbl.cell(0, nonel_col0 + n_nonel - 1))
            _write(m, 'NONEL', bold=True, size=7); _shd(m, HDR_BG); _vmid(m)

            run_start = nonel_col0
            cur_sub = nonel_meta[columns_subset[n_fixed][0]][0]
            for i in range(n_fixed, n_fixed + n_nonel + 1):
                cur_col = nonel_col0 + (i - n_fixed)
                yid_i = columns_subset[i][0] if i < n_fixed + n_nonel else None
                sub = nonel_meta[yid_i][0] if yid_i is not None else None
                if sub != cur_sub:
                    if cur_col - 1 > run_start:
                        m2 = tbl.cell(1, run_start).merge(tbl.cell(1, cur_col - 1))
                    else:
                        m2 = tbl.cell(1, run_start)
                    _write(m2, GROUPED_SUB_LABEL.get(cur_sub, cur_sub), bold=True, size=7)
                    _shd(m2, SUB_BG); _vmid(m2)
                    run_start = cur_col
                    cur_sub = sub

            code_counts = {}
            for i in range(n_fixed, n_fixed + n_nonel):
                sub, code, length = nonel_meta[columns_subset[i][0]]
                code_counts[(sub, code)] = code_counts.get((sub, code), 0) + 1
            for i in range(n_fixed, n_fixed + n_nonel):
                sub, code, length = nonel_meta[columns_subset[i][0]]
                col_i = nonel_col0 + (i - n_fixed)
                if sub == 'UNIDET':
                    leaf_txt = length
                elif code_counts[(sub, code)] > 1:
                    leaf_txt = f"{code}\n{length}"
                else:
                    leaf_txt = code
                cell = tbl.cell(2, col_i)
                _write(cell, leaf_txt, bold=True, size=7); _shd(cell, GRP_BG); _vmid(cell)
            col = nonel_col0 + n_nonel
        for lbl in tail_labels:
            m = tbl.cell(0, col).merge(tbl.cell(2, col))
            _write(m, lbl, bold=True, size=7); _shd(m, HDR_BG); _vmid(m)
            col += 1

    def build_agores_page(columns_subset, is_full):
        """is_full=True: πλήρης πίνακας (Α/Α, Αδεια, [υλικά], Ημερ.Αγ.,
        Προμηθευτής). is_full=False: συμπαγής πίνακας μόνο με Α/Α +
        Ημερ.Αγ. σαν άγκυρες (χρησιμοποιείται μόνο στα legacy nonel_mode
        όταν σπάει σε 2 πίνακες)."""
        n_mat = len(columns_subset)
        mat_w = material_widths_for(columns_subset)
        if is_full:
            lead_labels = ['Α/Α', 'Αρ.Άδ./\nΕκδ.Αρχή']
            tail_labels = ['Ημερ.Αγ./\nΑρ.Δελτ.', 'Στοιχεία\nΠρομηθευτή']
            widths = [0.8, 2.5] + mat_w + [3.0, 3.5]
        else:
            lead_labels = ['Α/Α', 'Ημερ.Αγ./\nΑρ.Δελτ.']
            tail_labels = []
            widths = [0.8, 3.0] + mat_w

        n_cols = len(widths)
        n_hdr_rows = 3 if grouped else 1
        tbl = doc.add_table(rows=n_hdr_rows, cols=n_cols)
        tbl.style = 'Table Grid'
        _setup_grid(tbl, widths)

        if grouped:
            build_group_header(tbl, columns_subset, lead_labels, tail_labels)
        else:
            hdrs = lead_labels + [f"{on}\n({mo})" for _, (on, mo) in columns_subset] + tail_labels
            for i, cell in enumerate(tbl.rows[0].cells):
                _shd(cell, HDR_BG)
                _write(cell, hdrs[i], bold=True, size=7)
                _tcW(cell, widths[i])
                _vmid(cell)

        n_fixed_pre = len(lead_labels)
        for ri, row in enumerate(rows):
            is_epi = row['type'] == 'epistrofi'
            dr = tbl.add_row()

            if is_full:
                adeia_ekd = row['adeia'] + (f"\n{row['ekdousa']}" if row.get('ekdousa') else '')
                data = ([str(row['aa']) if row['aa'] else '', adeia_ekd] +
                        [get_val(yid, row['ylika']) for yid, _ in columns_subset] +
                        [f"{fmt_date(row['imerominia'])}\n{row['parstatiko']}", row['promitheftis']])
            else:
                data = ([str(row['aa']) if row['aa'] else '',
                         f"{fmt_date(row['imerominia'])}\n{row['parstatiko']}"] +
                        [get_val(yid, row['ylika']) for yid, _ in columns_subset])

            fg = EPY_FG if is_epi else None
            bg = EPY_BG if is_epi else (ALT_BG if ri % 2 == 1 else None)

            for i, cell in enumerate(dr.cells):
                is_mat = n_fixed_pre <= i < n_fixed_pre + n_mat
                _write(cell, data[i], bold=is_epi, size=7, fg=fg,
                       align='right' if is_mat else 'center')
                if bg:
                    _shd(cell, bg)
                _tcW(cell, widths[i])
                _vmid(cell)

    def build_katanaloseis_page(columns_subset, is_full):
        """Ίδια δομή με build_agores_page αλλά με δεδομένα κατανάλωσης
        (kat_by_parst) — έλειπε εντελώς από το docx export."""
        n_mat = len(columns_subset)
        mat_w = material_widths_for(columns_subset)
        if is_full:
            lead_labels = ['Ημερ.\nΕισαγ.\nΑποθ.', 'Ημερ.\nΚατανάλ.']
            tail_labels = ['Παρατηρήσεις']
            widths = [1.6, 1.6] + mat_w + [3.5]
        else:
            lead_labels = ['Ημερ.\nΕισαγ.\nΑποθ.', 'Ημερ.\nΚατανάλ.']
            tail_labels = []
            widths = [1.6, 1.6] + mat_w

        n_cols = len(widths)
        n_hdr_rows = 3 if grouped else 1
        tbl = doc.add_table(rows=n_hdr_rows, cols=n_cols)
        tbl.style = 'Table Grid'
        _setup_grid(tbl, widths)

        if grouped:
            build_group_header(tbl, columns_subset, lead_labels, tail_labels)
        else:
            hdrs = lead_labels + [f"{on}\n({mo})" for _, (on, mo) in columns_subset] + tail_labels
            for i, cell in enumerate(tbl.rows[0].cells):
                _shd(cell, HDR_BG)
                _write(cell, hdrs[i], bold=True, size=7)
                _tcW(cell, widths[i])
                _vmid(cell)

        n_fixed_pre = len(lead_labels)
        kat_used = set()
        ri = 0
        for row in rows:
            is_epi = row['type'] == 'epistrofi'
            ap   = row.get('parstatiko', '')
            imer = row['imerominia']
            if is_epi or ap in kat_used:
                continue
            kat = kat_by_parst.get(ap, kat_by_parst.get(imer, {}))
            kat_used.add(ap)
            kat_imer  = kat.get('imerominia', imer)
            kat_ylika = kat.get('ylika', {})

            dr = tbl.add_row()
            if is_full:
                data = ([fmt_date(kat_imer), fmt_date(kat_imer)] +
                        [get_val(yid, kat_ylika) for yid, _ in columns_subset] +
                        [kat.get('paratirishis', '')])
            else:
                data = ([fmt_date(kat_imer), fmt_date(kat_imer)] +
                        [get_val(yid, kat_ylika) for yid, _ in columns_subset])

            bg = ALT_BG if ri % 2 == 1 else None
            for i, cell in enumerate(dr.cells):
                is_mat = n_fixed_pre <= i < n_fixed_pre + n_mat
                _write(cell, data[i], size=7, align='right' if is_mat else 'center')
                if bg:
                    _shd(cell, bg)
                _tcW(cell, widths[i])
                _vmid(cell)
            ri += 1

    if grouped:
        add_title('ΒΙΒΛΙΟ ΑΓΟΡΑΣ ΚΑΙ ΚΑΤΑΝΑΛΩΣΗΣ ΕΚΡΗΚΤΙΚΩΝ ΥΛΩΝ — ΑΓΟΡΕΣ / ΕΠΙΣΤΡΟΦΕΣ')
        build_agores_page(columns, is_full=True)
        doc.add_page_break()
        add_title('ΒΙΒΛΙΟ ΑΓΟΡΑΣ ΚΑΙ ΚΑΤΑΝΑΛΩΣΗΣ ΕΚΡΗΚΤΙΚΩΝ ΥΛΩΝ — ΚΑΤΑΝΑΛΩΣΕΙΣ')
        build_katanaloseis_page(columns, is_full=True)
    else:
        add_title('ΒΙΒΛΙΟ ΑΓΟΡΑΣ ΚΑΙ ΚΑΤΑΝΑΛΩΣΗΣ ΕΚΡΗΚΤΙΚΩΝ ΥΛΩΝ — ΑΓΟΡΕΣ / ΕΠΙΣΤΡΟΦΕΣ')
        build_agores_page(non_nonel, is_full=True)
        if nonel_cols:
            doc.add_page_break()
            add_title('ΒΙΒΛΙΟ ΑΓΟΡΑΣ ΚΑΙ ΚΑΤΑΝΑΛΩΣΗΣ ΕΚΡΗΚΤΙΚΩΝ ΥΛΩΝ — NONEL')
            build_agores_page(nonel_cols, is_full=False)

        doc.add_page_break()
        add_title('ΒΙΒΛΙΟ ΑΓΟΡΑΣ ΚΑΙ ΚΑΤΑΝΑΛΩΣΗΣ ΕΚΡΗΚΤΙΚΩΝ ΥΛΩΝ — ΚΑΤΑΝΑΛΩΣΕΙΣ')
        build_katanaloseis_page(non_nonel, is_full=True)
        if nonel_cols:
            doc.add_page_break()
            add_title('ΒΙΒΛΙΟ ΑΓΟΡΑΣ ΚΑΙ ΚΑΤΑΝΑΛΩΣΗΣ ΕΚΡΗΚΤΙΚΩΝ ΥΛΩΝ — NONEL (ΚΑΤΑΝΑΛΩΣΕΙΣ)')
            build_katanaloseis_page(nonel_cols, is_full=False)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ─── PDF ΥΠΟΛΟΓΙΣΜΟΥ ─────────────────────────────────────────────────────────

def export_ypologismos_pdf(parstatiko_agoras: str, senario: int, grammes: list) -> bytes:
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import cm
    from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT

    F, FB = register_fonts()
    if not FONT_REGULAR:
        F, FB = 'Helvetica', 'Helvetica-Bold'

    buf = __import__('io').BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4,
                            leftMargin=2*cm, rightMargin=2*cm,
                            topMargin=2*cm, bottomMargin=2*cm)

    TS = ParagraphStyle('t', fontSize=13, fontName=FB, alignment=TA_CENTER, spaceAfter=4)
    SS = ParagraphStyle('s', fontSize=9,  fontName=F,  alignment=TA_CENTER, spaceAfter=12)
    HS = ParagraphStyle('h', fontSize=9,  fontName=FB, alignment=TA_CENTER, leading=11)
    CS = ParagraphStyle('c', fontSize=9,  fontName=F,  alignment=TA_CENTER, leading=11)
    AS = ParagraphStyle('a', fontSize=10, fontName=FB, alignment=TA_CENTER, leading=12,
                        textColor=colors.HexColor('#2d6a4f'))

    from datetime import datetime
    senario_txt = 'Σενάριο 1: Αγορά + Κατανάλωση → Επιστροφή' if senario==1 \
                  else 'Σενάριο 2: Αγορά + Επιστροφή → Κατανάλωση'

    story = [
        Paragraph("ΔΕΛΤΙΟ ΥΠΟΛΟΓΙΣΜΟΥ ΕΚΡΗΚΤΙΚΩΝ ΥΛΩΝ", TS),
        Paragraph(f"Παραστατικό Αγοράς: {parstatiko_agoras}  |  {senario_txt}", SS),
        Paragraph(f"Ημερομηνία Εκτύπωσης: {datetime.now().strftime('%d/%m/%Y %H:%M')}", SS),
    ]

    input_lbl  = 'Κατανάλωση' if senario==1 else 'Επιστροφή'
    result_lbl = 'Επιστροφή'  if senario==1 else 'Κατανάλωση'

    headers = ['Υλικό', 'Μονάδα', 'Αγορά', input_lbl, result_lbl]
    rows = [headers]
    for g in grammes:
        val = g['posotita_epistrofis'] if senario==1 else g['posotita_katanalosis']
        inp = g['posotita_katanalosis'] if senario==1 else g['posotita_epistrofis']
        rows.append([
            g['yliko_onoma'], g['monada'],
            f"{g['posotita_agoras']:,.3f}".replace(',','X').replace('.',',').replace('X','.'),
            f"{inp:,.3f}".replace(',','X').replace('.',',').replace('X','.'),
            f"{val:,.3f}".replace(',','X').replace('.',',').replace('X','.')
        ])

    col_w = [7*cm, 1.5*cm, 2.5*cm, 2.5*cm, 2.5*cm]
    t = Table(rows, colWidths=col_w)
    t.setStyle(TableStyle([
        ('BACKGROUND',   (0,0), (-1,0), colors.HexColor('#1a365d')),
        ('TEXTCOLOR',    (0,0), (-1,0), colors.white),
        ('FONTNAME',     (0,0), (-1,0), FB),
        ('FONTNAME',     (0,1), (-1,-1), F),
        ('FONTSIZE',     (0,0), (-1,-1), 9),
        ('ALIGN',        (0,0), (-1,-1), 'CENTER'),
        ('ALIGN',        (0,1), (0,-1), 'LEFT'),
        ('BACKGROUND',   (4,1), (4,-1), colors.HexColor('#e6f7ee')),
        ('FONTNAME',     (4,1), (4,-1), FB),
        ('ROWBACKGROUNDS',(0,1),(-1,-1),[colors.white, colors.HexColor('#f7f9fc')]),
        ('GRID',         (0,0), (-1,-1), 0.5, colors.HexColor('#cbd5e0')),
        ('TOPPADDING',   (0,0), (-1,-1), 5),
        ('BOTTOMPADDING',(0,0), (-1,-1), 5),
    ]))
    story.append(t)
    story.append(Spacer(1, 1*cm))
    story.append(Paragraph(f"Αποτέλεσμα {result_lbl} — Για χρήση κατά την έκδοση παραστατικού", AS))

    doc.build(story)
    return buf.getvalue()


def export_lista_agores(kiniseis: list, apo_label: str, eos_label: str) -> bytes:
    """Κατάσταση Αγορών: στήλες=υλικά με grouped headers, γραμμές=αγορές, υποσύνολο ανά άδεια."""
    import openpyxl
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    from openpyxl.utils import get_column_letter
    from io import BytesIO

    agores_kin = [k for k in kiniseis if k['tipos'] == 'ΕΙΣΑΓΩΓΗ']

    FONT = 'Iosevka'

    # ── Σειρά ομάδων υλικών ───────────────────────────────────────────────────────
    GROUP_ORDER = [None, 'POLADYN', 'ΘΡΥΑΛΛΙΔΑ', 'ΠΥΡΟΚΡΟΤΗΤΕΣ',
                   'NONEL SNAPLINE', 'NONEL LP', 'NONEL UNIDET']
    GROUP_COLORS = {
        None:             "4A7FC1",
        'POLADYN':        "4A7FC1",
        'ΘΡΥΑΛΛΙΔΑ':     "4A7FC1",
        'ΠΥΡΟΚΡΟΤΗΤΕΣ':  "4A7FC1",
        'NONEL SNAPLINE': "1A365D",
        'NONEL LP':       "1A365D",
        'NONEL UNIDET':   "1A365D",
    }
    GROUP_SUB_COLORS = {
        'POLADYN':        "2C5282",
        'ΘΡΥΑΛΛΙΔΑ':     "2C5282",
        'ΠΥΡΟΚΡΟΤΗΤΕΣ':  "2C5282",
        'NONEL SNAPLINE': "2C5282",
        'NONEL LP':       "2C5282",
        'NONEL UNIDET':   "2C5282",
    }

    # ── Λίστα υλικών με σειρά ─────────────────────────────────────────────────────
    yliko_data = {}   # yid -> {onoma, monada, export_group}
    yid_order  = []
    for k in agores_kin:
        yid = k['yliko_id']
        if yid not in yliko_data:
            yliko_data[yid] = {
                'onoma':        k['yliko_onoma'],
                'monada':       k['monada_metrisis'],
                'export_group': k.get('export_group'),
            }
            yid_order.append(yid)

    buckets = {g: [] for g in GROUP_ORDER}
    extra   = []
    for yid in yid_order:
        grp = yliko_data[yid]['export_group']
        if grp in buckets:
            buckets[grp].append(yid)
        else:
            extra.append(yid)
    ordered_yids = []
    for g in GROUP_ORDER:
        ordered_yids.extend(buckets[g])
    ordered_yids.extend(extra)

    vo = [(yid, yliko_data[yid]) for yid in ordered_yids]
    n  = len(vo)

    # Ομαδοποίηση συνεχόμενων στηλών ανά export_group
    col_groups = []  # list of (grp, [col_idx_0based, ...])
    for i, (yid, yd) in enumerate(vo):
        grp = yd['export_group']
        if col_groups and col_groups[-1][0] == grp:
            col_groups[-1][1].append(i)
        else:
            col_groups.append((grp, [i]))

    # ── Ομαδοποίηση κινήσεων σε γραμμές αγοράς ───────────────────────────────────
    rows_dict = OrderedDict()
    for k in agores_kin:
        yid   = k['yliko_id']
        imer  = k['imerominia']
        parst = k.get('arithmos_parstatikos') or ''
        adeia = k.get('arithmos_adeias') or ''
        prom  = k.get('promitheftis_syntomografia') or k.get('promitheftis_onoma') or ''
        key   = (imer, parst, adeia, prom)
        if key not in rows_dict:
            rows_dict[key] = {'imerominia': imer, 'parstatiko': parst,
                              'adeia': adeia, 'promitheftis': prom, 'ylika': {}}
        rows_dict[key]['ylika'][yid] = rows_dict[key]['ylika'].get(yid, 0) + k['posotita']

    rows = sorted(rows_dict.values(), key=lambda r: (r['imerominia'], r['parstatiko']))

    # Ομαδοποίηση ανά άδεια (consecutive)
    groups = []
    cur_adeia, cur_group = None, []
    for row in rows:
        if row['adeia'] != cur_adeia:
            if cur_group:
                groups.append((cur_adeia, cur_group))
            cur_adeia, cur_group = row['adeia'], [row]
        else:
            cur_group.append(row)
    if cur_group:
        groups.append((cur_adeia, cur_group))

    # ── Styles ───────────────────────────────────────────────────────────────────
    thin      = Side(style='thin', color="AABBD0")
    thick_bot = Side(style='medium', color="4A7FC1")
    brd       = Border(left=thin, right=thin, top=thin, bottom=thin)
    brd_sub   = Border(left=thin, right=thin, top=thick_bot, bottom=thick_bot)

    NAVY  = "1A365D"; LNAV = "4A7FC1"; LBLU = "BEE3F8"
    ADEIA = "2C5282"; GRAND = "1A365D"; ALT  = "EEF2F7"

    def hdr_style(cell, bg, bold=True, sz=9, wrap=False, halign='center', fg="FFFFFF"):
        cell.fill = PatternFill("solid", fgColor=bg)
        cell.font = Font(name=FONT, bold=bold, color=fg, size=sz)
        cell.alignment = Alignment(horizontal=halign, vertical='center', wrap_text=wrap)
        cell.border = brd

    def data_style(cell, bg=None, bold=False, color="000000", halign='left', num=False):
        if bg:
            cell.fill = PatternFill("solid", fgColor=bg)
        cell.font = Font(name=FONT, bold=bold, color=color, size=9)
        cell.alignment = Alignment(horizontal='right' if num else halign,
                                   vertical='center', wrap_text=False)
        cell.border = brd

    # ── Workbook ──────────────────────────────────────────────────────────────────
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = 'Κατάσταση Αγορών'

    TOTAL_COLS = 4 + n  # Α/Α | Ημ/νία | Παραστατικό | Προμηθευτής | [υλικά]
    TC = get_column_letter(TOTAL_COLS)

    # ── Row 1: Τίτλος ─────────────────────────────────────────────────────────────
    ws.merge_cells(f'A1:{TC}1')
    ws['A1'] = f'ΚΑΤΑΣΤΑΣΗ ΑΓΟΡΩΝ — {apo_label} έως {eos_label}'
    ws['A1'].fill = PatternFill("solid", fgColor=NAVY)
    ws['A1'].font = Font(name=FONT, bold=True, color="FFFFFF", size=12)
    ws['A1'].alignment = Alignment(horizontal='center', vertical='center')
    ws.row_dimensions[1].height = 22

    # ── Row 2: Group headers + fixed headers ──────────────────────────────────────
    # Fixed columns: merge rows 2-3
    fixed_hdrs = ['Α/Α', 'Ημ/νία', 'Παραστατικό', 'Προμηθευτής']
    widths     = [5, 11, 18, 18]
    for ci, (h, w) in enumerate(zip(fixed_hdrs, widths), 1):
        ws.merge_cells(start_row=2, start_column=ci, end_row=3, end_column=ci)
        hdr_style(ws.cell(row=2, column=ci, value=h), LNAV, wrap=True)
        ws.column_dimensions[get_column_letter(ci)].width = w

    # Material columns grouped
    for grp, idxs in col_groups:
        col_start = 5 + idxs[0]
        col_end   = 5 + idxs[-1]
        grp_color = GROUP_COLORS.get(grp, LNAV)
        sub_color = GROUP_SUB_COLORS.get(grp, LNAV)

        if grp is None:
            # Standalone: merge rows 2-3
            for idx in idxs:
                ci  = 5 + idx
                yd  = vo[idx][1]
                lbl = f"{yd['onoma']}\n({yd['monada']})"
                ws.merge_cells(start_row=2, start_column=ci, end_row=3, end_column=ci)
                hdr_style(ws.cell(row=2, column=ci, value=lbl), grp_color, wrap=True)
                ws.column_dimensions[get_column_letter(ci)].width = 13
        else:
            # Row 2: group name merged across all sub-columns
            if len(idxs) > 1:
                ws.merge_cells(start_row=2, start_column=col_start, end_row=2, end_column=col_end)
            hdr_style(ws.cell(row=2, column=col_start, value=grp), grp_color, sz=10)
            # Row 3: individual material names
            for idx in idxs:
                ci  = 5 + idx
                yd  = vo[idx][1]
                lbl = f"{yd['onoma']}\n({yd['monada']})"
                hdr_style(ws.cell(row=3, column=ci, value=lbl), sub_color, sz=8, wrap=True)
                ws.column_dimensions[get_column_letter(ci)].width = 13

    ws.row_dimensions[2].height = 20
    ws.row_dimensions[3].height = 36

    # ── Data (from row 4) ─────────────────────────────────────────────────────────
    ri    = 4
    aa    = 0
    grand = {yid: 0.0 for yid, _ in vo}

    for adeia_num, adeia_rows in groups:
        adeia_lbl = adeia_num if adeia_num else '(χωρίς άδεια)'

        # Άδεια header
        ws.merge_cells(f'A{ri}:{TC}{ri}')
        cell = ws.cell(row=ri, column=1, value=f'Άδεια: {adeia_lbl}')
        cell.fill = PatternFill("solid", fgColor=ADEIA)
        cell.font = Font(name=FONT, bold=True, color="FFFFFF", size=9)
        cell.alignment = Alignment(horizontal='left', vertical='center', indent=1)
        cell.border = brd
        ws.row_dimensions[ri].height = 16
        ri += 1

        adeia_sums = {yid: 0.0 for yid, _ in vo}
        alt = False

        for row in adeia_rows:
            aa += 1
            bg  = ALT if alt else None
            alt = not alt

            c1 = ws.cell(row=ri, column=1, value=aa)
            c1.border = brd
            c1.font = Font(name=FONT, size=9)
            c1.alignment = Alignment(horizontal='center', vertical='center')
            if bg: c1.fill = PatternFill("solid", fgColor=bg)

            for ci, val in enumerate([fmt_date(row['imerominia']), row['parstatiko'], row['promitheftis']], 2):
                data_style(ws.cell(row=ri, column=ci, value=val), bg)

            for col_idx, (yid, _) in enumerate(vo, 5):
                v    = row['ylika'].get(yid) or None
                cell = ws.cell(row=ri, column=col_idx, value=fmt_num(v) if v else '')
                data_style(cell, bg, num=True)
                if v:
                    adeia_sums[yid] += v
                    grand[yid]      += v

            ws.row_dimensions[ri].height = 14
            ri += 1

        # Υποσύνολο άδειας
        ws.merge_cells(f'A{ri}:D{ri}')
        cell = ws.cell(row=ri, column=1, value=f'Υποσύνολο Άδεια {adeia_lbl}')
        cell.fill = PatternFill("solid", fgColor=LBLU)
        cell.font = Font(name=FONT, bold=True, color=NAVY, size=9)
        cell.alignment = Alignment(horizontal='right', vertical='center')
        cell.border = brd_sub
        for ci in range(2, 5):
            ws.cell(row=ri, column=ci).fill   = PatternFill("solid", fgColor=LBLU)
            ws.cell(row=ri, column=ci).border = brd_sub

        for col_idx, (yid, _) in enumerate(vo, 5):
            v    = adeia_sums[yid] or None
            cell = ws.cell(row=ri, column=col_idx, value=fmt_num(v) if v else '')
            cell.fill      = PatternFill("solid", fgColor=LBLU)
            cell.font      = Font(name=FONT, bold=True, color=NAVY, size=9)
            cell.alignment = Alignment(horizontal='right', vertical='center')
            cell.border    = brd_sub

        ws.row_dimensions[ri].height = 15
        ri += 1

    # Γενικό σύνολο
    ws.merge_cells(f'A{ri}:D{ri}')
    cell = ws.cell(row=ri, column=1, value='ΓΕΝΙΚΟ ΣΥΝΟΛΟ')
    cell.fill      = PatternFill("solid", fgColor=GRAND)
    cell.font      = Font(name=FONT, bold=True, color="FFFFFF", size=10)
    cell.alignment = Alignment(horizontal='right', vertical='center')
    cell.border    = brd
    for ci in range(2, 5):
        ws.cell(row=ri, column=ci).fill   = PatternFill("solid", fgColor=GRAND)
        ws.cell(row=ri, column=ci).border = brd

    for col_idx, (yid, _) in enumerate(vo, 5):
        v    = grand[yid] or None
        cell = ws.cell(row=ri, column=col_idx, value=fmt_num(v) if v else '')
        cell.fill      = PatternFill("solid", fgColor=GRAND)
        cell.font      = Font(name=FONT, bold=True, color="FFFFFF", size=10)
        cell.alignment = Alignment(horizontal='right', vertical='center')
        cell.border    = brd

    ws.row_dimensions[ri].height = 18

    ws.freeze_panes = 'A4'

    # Page setup — landscape A3, fit to width, margins από reference
    ws.page_setup.orientation = 'landscape'
    ws.page_setup.paperSize   = 8   # A3
    ws.page_setup.scale       = 85
    ws.page_setup.fitToWidth  = 1
    ws.page_setup.fitToHeight = 0
    ws.page_margins.left   = 0.354167
    ws.page_margins.right  = 0.196528
    ws.page_margins.top    = 0.236111
    ws.page_margins.bottom = 0.236111
    ws.page_margins.header = 0
    ws.page_margins.footer = 0

    buf = BytesIO()
    wb.save(buf)
    return buf.getvalue()


# ── ΔΕΛΤΙΟ ΔΡΑΣΤΗΡΙΟΤΗΤΑΣ ─────────────────────────────────────────────────────
# Σύνολα κατανάλωσης ανά νόμιμη κατηγορία εκρηκτικού, για μια περίοδο.

def _deltio_sums(kiniseis):
    """Κατανάλωση ανά κατηγορία = ΑΓΟΡΑ − ΕΠΙΣΤΡΟΦΗ (όπως στον Υπολογιστή, Σενάριο 2),
    αφού η κατανάλωση δεν καταχωρείται πάντα ως ξεχωριστή κίνηση."""
    from database import NOMIKES_KATIGORIES, MONADES_NOMIKON_KATIGORION
    agores     = OrderedDict((kat, 0.0) for kat in NOMIKES_KATIGORIES)
    epistrofes = OrderedDict((kat, 0.0) for kat in NOMIKES_KATIGORIES)
    for k in kiniseis:
        kat = k.get('nomiki_katigoria')
        if kat not in agores:
            continue
        if k['tipos'] == 'ΕΙΣΑΓΩΓΗ':
            agores[kat] += k['posotita']
        elif k['tipos'] == 'ΕΠΙΣΤΡΟΦΗ':
            epistrofes[kat] += k['posotita']
    return OrderedDict((kat, max(0.0, agores[kat] - epistrofes[kat])) for kat in NOMIKES_KATIGORIES)

def export_deltio_drastiriotitas_excel(kiniseis: list, apo_label: str, eos_label: str, adeia_label: str = None) -> bytes:
    import openpyxl
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    from io import BytesIO
    from database import MONADES_NOMIKON_KATIGORION

    sums = _deltio_sums(kiniseis)

    FONT = 'Iosevka'
    NAVY = "1A365D"; LBLU = "BEE3F8"; ALT = "EEF2F7"
    thin = Side(style='thin', color="AABBD0")
    brd  = Border(left=thin, right=thin, top=thin, bottom=thin)

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = 'Δελτίο Δραστηριότητας'

    ws.merge_cells('A1:C1')
    c = ws.cell(row=1, column=1, value='ΔΕΛΤΙΟ ΔΡΑΣΤΗΡΙΟΤΗΤΑΣ — ΣΥΝΟΛΑ ΚΑΤΑΝΑΛΩΣΗΣ ΑΝΑ ΚΑΤΗΓΟΡΙΑ')
    c.font = Font(name=FONT, bold=True, size=13)
    c.alignment = Alignment(horizontal='center')

    period_text = f'Περίοδος: {apo_label} έως {eos_label}' if (apo_label or eos_label) else 'Περίοδος: Όλο το χρονικό διάστημα'
    subtitle = f'Άδεια: {adeia_label} — {period_text}' if adeia_label else period_text
    ws.merge_cells('A2:C2')
    c = ws.cell(row=2, column=1, value=subtitle)
    c.font = Font(name=FONT, size=10)
    c.alignment = Alignment(horizontal='center')

    headers = ['Κατηγορία Εκρηκτικού', 'Μονάδα', 'Κατανάλωση']
    for ci, h in enumerate(headers, 1):
        cell = ws.cell(row=4, column=ci, value=h)
        cell.fill      = PatternFill("solid", fgColor=NAVY)
        cell.font      = Font(name=FONT, bold=True, color="FFFFFF", size=10)
        cell.alignment = Alignment(horizontal='center', vertical='center')
        cell.border    = brd

    ri = 5
    for kat, total in sums.items():
        cell = ws.cell(row=ri, column=1, value=kat)
        cell.font = Font(name=FONT, size=10)
        cell.alignment = Alignment(horizontal='left')
        cell.border = brd
        if ri % 2 == 0:
            cell.fill = PatternFill("solid", fgColor=ALT)

        cell = ws.cell(row=ri, column=2, value=MONADES_NOMIKON_KATIGORION[kat])
        cell.font = Font(name=FONT, size=10)
        cell.alignment = Alignment(horizontal='center')
        cell.border = brd
        if ri % 2 == 0:
            cell.fill = PatternFill("solid", fgColor=ALT)

        cell = ws.cell(row=ri, column=3, value=fmt_num(total) or '0,00')
        cell.font = Font(name=FONT, size=10)
        cell.alignment = Alignment(horizontal='right')
        cell.border = brd
        if ri % 2 == 0:
            cell.fill = PatternFill("solid", fgColor=ALT)
        ri += 1

    ws.column_dimensions['A'].width = 32
    ws.column_dimensions['B'].width = 10
    ws.column_dimensions['C'].width = 16

    ws.page_setup.orientation = 'portrait'
    ws.page_setup.fitToWidth  = 1
    ws.page_setup.fitToHeight = 0

    buf = BytesIO()
    wb.save(buf)
    return buf.getvalue()

def export_deltio_drastiriotitas_pdf(kiniseis: list, apo_label: str, eos_label: str, adeia_label: str = None) -> bytes:
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import cm
    from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.enums import TA_CENTER
    from io import BytesIO
    from database import MONADES_NOMIKON_KATIGORION

    fname, fbold = register_fonts()
    F  = fname if FONT_REGULAR else 'Helvetica'
    FB = fbold if FONT_BOLD    else 'Helvetica-Bold'

    sums = _deltio_sums(kiniseis)

    buf = BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4,
                            leftMargin=2*cm, rightMargin=2*cm,
                            topMargin=2*cm, bottomMargin=2*cm)

    TS = ParagraphStyle('t', fontSize=13, fontName=FB, alignment=TA_CENTER, spaceAfter=4)
    SS = ParagraphStyle('s', fontSize=9,  fontName=F,  alignment=TA_CENTER, spaceAfter=12)

    period_text = f'Περίοδος: {apo_label} έως {eos_label}' if (apo_label or eos_label) else 'Περίοδος: Όλο το χρονικό διάστημα'
    subtitle = f'Άδεια: {adeia_label} — {period_text}' if adeia_label else period_text
    story = [
        Paragraph("ΔΕΛΤΙΟ ΔΡΑΣΤΗΡΙΟΤΗΤΑΣ — ΣΥΝΟΛΑ ΚΑΤΑΝΑΛΩΣΗΣ ΑΝΑ ΚΑΤΗΓΟΡΙΑ", TS),
        Paragraph(subtitle, SS),
    ]

    rows = [['Κατηγορία Εκρηκτικού', 'Μονάδα', 'Κατανάλωση']]
    for kat, total in sums.items():
        rows.append([kat, MONADES_NOMIKON_KATIGORION[kat], fmt_num(total) or '0,00'])

    col_w = [9*cm, 2.5*cm, 3.5*cm]
    t = Table(rows, colWidths=col_w)
    t.setStyle(TableStyle([
        ('BACKGROUND',    (0,0), (-1,0), colors.HexColor('#1a365d')),
        ('TEXTCOLOR',     (0,0), (-1,0), colors.white),
        ('FONTNAME',      (0,0), (-1,0), FB),
        ('FONTNAME',      (0,1), (-1,-1), F),
        ('FONTSIZE',      (0,0), (-1,-1), 10),
        ('ALIGN',         (0,0), (-1,-1), 'CENTER'),
        ('ALIGN',         (0,1), (0,-1), 'LEFT'),
        ('ROWBACKGROUNDS',(0,1), (-1,-1), [colors.white, colors.HexColor('#f7f9fc')]),
        ('GRID',          (0,0), (-1,-1), 0.5, colors.HexColor('#cbd5e0')),
        ('TOPPADDING',    (0,0), (-1,-1), 6),
        ('BOTTOMPADDING', (0,0), (-1,-1), 6),
    ]))
    story.append(t)

    doc.build(story)
    return buf.getvalue()
